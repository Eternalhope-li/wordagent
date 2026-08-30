"""编辑模式：根据自然语言指令修改已有 docx。

安全与质量设计（重要文件）：
- 自动备份原文件到 output/backups/；默认另存为新文件
- 目标匹配三级策略：精确 -> 包含 -> 近似（模糊）；仍有歧义时让用户选择
- 所有操作先全量校验，任一失败整体中止，不写盘
- 落笔支持 `##/###` 标题、列表等标记
- 程序化结构校验：原标题完整性、操作生效性（不依赖 LLM）
- AI 复核为【建议模式】：只报告问题，不自动落笔（除非显式 auto_fix）
- 保存后重新打开校验完整性
"""
from __future__ import annotations

import difflib
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

from .config import Config
from .extractor import docx_to_markdown
from .llm import LLMClient, LLMError
from .memory import Memory
from .renderer import INLINE_RE

PLAN_PROMPT = """你是一名严谨的文档编辑助手。用户会用自然语言提出对现有 Word 文档的修改要求。

你要把修改翻译成一组【原子操作】，严格只输出 JSON，不要输出任何其他文字。JSON 格式：
{
  "summary": "一句话说明本次修改",
  "operations": [
    {"op": "replace", "target": "原文中要替换的段落文本", "new_text": "替换后的新内容"},
    {"op": "insert_after", "target": "原文中某个段落文本", "new_text": "要插入的内容"},
    {"op": "delete", "target": "原文中要删除的段落文本"},
    {"op": "set_style", "target": "原文中某个段落文本", "style": "Heading 1|Heading 2|Heading 3|Normal|List Bullet|List Number"}
  ]
}

操作说明：
- replace：整段替换。new_text 支持多行，可用 ##/### 标题、**加粗**、*斜体*、`代码`、- 无序列表、1. 有序列表。
- insert_after：在目标段落后插入新内容（可多行）。
- delete：删除目标段落。
- set_style：把目标段落改为指定样式。
- insert_after 的目标也可以是表格行（形如 `| 单元格 | 单元格 |`），系统会把内容插入到该表格之后。
- replace 的目标也可以是表格行（形如 `| a | b |`），new_text 用换行分隔各列新内容（如 `| 新值1 | 新值2 |` 对应两列），系统会按顺序替换该行单元格文本。
- delete 的目标也可以是表格行（形如 `| a | b |`），系统会删除该行。
- 新增章节（new_text 以 ##/### 开头）时，优先锚定到【章节标题】段落（##/### 标题行），不要锚定正文段落或表格行；系统会自动把新章节放到该章节末尾，避免打断子节结构。
- 如果【文档内要求】中写明了格式、字数、数据安排等约束（如批注里的实验安排、字体字号要求），修改时必须满足这些约束。

关于 target（铁律）：
1. target 必须来自【当前文档内容】，优先使用完整句子或足够长的唯一片段；禁止编造原文。
2. 文本里的 **加粗** / *斜体* / `代码` 标记请忽略，只抄写纯文字（系统会自动做容错匹配）。
3. 若用户要求涉及多处，请拆成多个 operation。
4. 只做用户明确要求的修改；其余内容一律原样保留，禁止重写全文、禁止改标题层级、禁止添加无关内容。
5. 新增章节/内容优先使用 insert_after 锚定在合适位置；新增章节（##/###）请锚定章节标题而非正文段落或表格行；若必须 replace 一个标题，请在 new_text 末尾用 `##`/`###` 标记把原标题保留回来。
6. 若用户要求整体重写文档，应提示建议使用"生成新文档"功能，而不是大规模 replace。
7. 内容充实度：替换/插入的正文要完整充实——实验步骤、操作过程类不少于 4 步，写明具体操作、参数、观察结果；分析/总结/说明类一般 80~200 字，覆盖要点。不要只写一两句或半截内容。
"""

CRITIC_PROMPT = """你是一名严格的文档质量检查员。请核对以下三点：
1) 【用户修改要求】是否已被【已执行的操作】满足；
2) 对比【修改前标题】与【修改后标题】：除用户明确要求删除/修改外，原有章节标题不得消失、不得变成正文、不得被其他标题覆盖；
3) 在【改动区域上下文】中确认修改真实生效（替换的新内容存在、删除的内容已不存在、插入的内容已出现）；
4) 检查替换/新增的内容是否完整充实：若用户要求的是实验步骤/操作流程，应有完整编号步骤（一般不少于 4 步）；若是分析/总结/说明，应覆盖要点、内容充实（一般 80~200 字）。内容过短、半截未写完属于未满足要求。

严格只输出 JSON：
{"satisfied": true, "reason": "简短说明", "issues": [], "operations": []}

规则：
- satisfied=true 表示要求已满足或要求本身较模糊、现有内容可接受，此时 operations 必须为空数组。
- satisfied=false 时，用 reason/issues 说明问题，operations 给出可执行的补充操作（格式与编辑操作一致：replace/insert_after/delete/set_style）。
- 只针对明确的问题给出少量操作；不要对已满足的内容做多余修改，严禁重写全文。
- 若文档较长、无法确认某处内容，视为可接受，不要臆断缺失。
"""

VALID_STYLES = {
    "Heading 1", "Heading 2", "Heading 3", "Heading 4",
    "Normal", "List Bullet", "List Number", "Quote",
}

OP_ALIASES = {
    "replace": "replace", "替换": "replace", "修改": "replace", "改写": "replace", "改": "replace",
    "insert_after": "insert_after", "插入": "insert_after", "insert": "insert_after", "追加": "insert_after",
    "delete": "delete", "删除": "delete",
    "set_style": "set_style", "样式": "set_style",
}

BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
NUMBER_RE = re.compile(r"^\s*\d+[.、)]\s+(.*)$")
HEADING_LINE_RE = re.compile(r"^(#{1,4})\s+(.*)$")
MD_MARKER_RE = re.compile(r"\*\*|\*|`")


class EditValidationError(RuntimeError):
    """编辑校验失败（目标未找到等），不会写盘。"""


class AmbiguousTargetError(EditValidationError):
    """目标匹配到多个段落，需要用户选择。"""

    def __init__(self, message: str, candidates: list[str]):
        super().__init__(message)
        self.candidates = candidates


# ---------------------------------------------------------------- 规划
def _smart_truncate(text: str, max_chars: int = 20000, per_section: int = 1200) -> tuple[str, bool]:
    """按章节结构智能截断：所有标题尽量保留，每节正文最多 per_section 字符。

    相比硬截断，模型能看见完整章节骨架，定位目标更准确（"找不到匹配文本"更少）。
    返回 (截断后文本, 是否截断)。
    """
    lines = text.splitlines()
    out: list[str] = []
    used = 0
    truncated = False
    section_len = 0
    for line in lines:
        stripped = line.strip()
        # 结构行（标题 / 表格行 / 分隔行）完整保留，不参与正文截断：
        # 表格内容往往就是用户要求修改的对象，截掉会导致 AI 看不到、改不动
        is_structural = (
            stripped.startswith("#")
            or stripped.startswith("|")
            or stripped == "---"
        )
        add = line + "\n"
        if is_structural:
            section_len = 0
            out.append(add)
            used += len(add)
            continue
        if section_len >= per_section:
            if not truncated:
                out.append("……（本节内容过长，已省略中间部分）……\n")
                truncated = True
                used += len("……（本节内容过长，已省略中间部分）……\n")
            continue
        out.append(add)
        used += len(add)
        section_len += len(add)
    # 章节过多导致整体仍超预算：从尾部删段落直至达标（标题最后删）
    while used > max_chars and len(out) > 6:
        removed = out.pop()
        used -= len(removed)
        truncated = True
    return "".join(out), truncated


def plan_edit(
    instruction: str,
    markdown: str,
    memory_context: str,
    llm: LLMClient,
    max_chars: int = 20000,
    doc_requirements: str = "",
) -> dict[str, Any]:
    content, truncated = _smart_truncate(markdown, max_chars)
    if truncated:
        content += "\n\n……（文档过长，已按章节结构精简展示，标题完整；定位时请使用原文片段）……"
    extra = ""
    if doc_requirements:
        extra = (
            "\n\n【文档内要求（批注/页眉页脚/文本框，必须遵守）】\n"
            + doc_requirements[:4000]
            + "\n（以上是文档里写的要求，如格式、字数、数据安排等，修改时必须满足）"
        )
    messages = [
        {"role": "system", "content": PLAN_PROMPT},
        {
            "role": "user",
            "content": (
                f"【历史记忆】\n{memory_context}\n\n"
                f"【用户修改要求】\n{instruction}\n\n"
                f"【当前文档内容】\n{content}"
                f"{extra}"
            ),
        },
    ]
    data = llm.chat_json(messages, temperature=0.2, max_tokens=8192)
    operations = _normalize_ops(data.get("operations"))
    if not operations:
        raise EditValidationError("模型未给出有效的修改操作，请重试或换一种说法。")
    return {
        "summary": str(data.get("summary", "")).strip(),
        "operations": operations,
        "truncated": truncated,
    }


def _normalize_ops(raw: Any) -> list[dict]:
    ops: list[dict] = []
    if not isinstance(raw, list):
        return ops
    for item in raw:
        if not isinstance(item, dict):
            continue
        kind = OP_ALIASES.get(str(item.get("op", "")).strip().lower(), "")
        target = str(item.get("target", "")).strip()
        if not kind or not target:
            continue
        op: dict[str, Any] = {"op": kind, "target": target}
        if "new_text" in item and item["new_text"] is not None:
            op["new_text"] = str(item["new_text"])
        if item.get("style"):
            op["style"] = str(item["style"])
        ops.append(op)
    return ops


# ---------------------------------------------------------------- 定位（三级匹配）
def _normalize_text(text: str) -> str:
    """去掉 markdown 标记、压缩空白，用于容错匹配。"""
    t = MD_MARKER_RE.sub("", text)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def _iter_all_paragraphs(doc: Document):
    """遍历正文与表格内的所有段落。"""
    body = doc.element.body

    def walk(element):
        for child in element.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, doc)
            elif child.tag == qn("w:tbl"):
                yield from walk(child)

    yield from walk(body)


def _find_paragraph(doc: Document, target: str) -> tuple[Paragraph, str]:
    raw = target.strip()
    norm = _normalize_text(raw)
    if not norm:
        raise EditValidationError("操作缺少有效的 target。")
    all_paras = [p for p in _iter_all_paragraphs(doc) if p.text.strip()]
    scored = [(_normalize_text(p.text), p) for p in all_paras]

    # 1) 精确匹配（归一化后）
    exact = [p for t, p in scored if t == norm]
    if len(exact) == 1:
        return exact[0], "精确匹配"
    if len(exact) > 1:
        raise AmbiguousTargetError(
            f"目标「{raw[:40]}」匹配到 {len(exact)} 个相同段落，请选择：", [p.text for p in exact]
        )

    # 2) 包含匹配
    contains = [p for t, p in scored if norm in t]
    if len(contains) == 1:
        return contains[0], "包含匹配"
    if len(contains) > 1:
        contains.sort(key=lambda p: len(_normalize_text(p.text)))
        if len(_normalize_text(contains[0].text)) < len(_normalize_text(contains[1].text)):
            return contains[0], "包含匹配（取最短）"
        raise AmbiguousTargetError(
            f"目标「{raw[:40]}」匹配到 {len(contains)} 个段落，请选择：",
            [p.text for p in contains[:5]],
        )

    # 3) 近似匹配（相似度最高且明显领先）
    ranked = sorted(
        ((difflib.SequenceMatcher(None, norm, t).ratio(), t, p) for t, p in scored),
        key=lambda x: -x[0],
    )
    best_ratio, _, best_para = ranked[0]
    second_ratio = ranked[1][0] if len(ranked) > 1 else 0.0
    if best_ratio >= 0.78 and best_ratio - second_ratio >= 0.05:
        return best_para, f"近似匹配（{best_ratio:.0%}）"

    hint = "\n".join(f"   · {t[:50]}" for t, _ in scored[:3])
    raise EditValidationError(
        f"未找到与「{raw[:50]}」匹配的内容（最接近的段落：\n{hint}）。\n"
        "安全保护已中止本次修改；请换用更接近原文的措辞重试。"
    )


def _find_target(doc: Document, target: str):
    """定位目标，返回 (kind, obj, how)。

    kind='paragraph'：obj 为 Paragraph；kind='table_row'：obj 为 (Table, row_index)。
    """
    try:
        para, how = _find_paragraph(doc, target)
        return "paragraph", para, how
    except AmbiguousTargetError:
        raise
    except EditValidationError:
        pass  # 段落无匹配，试表格行

    # 表格行匹配：提取格式为 `| a | b |`，单元格拼接为 `a | b`，统一去掉管道符后比对
    def norm_row(t: str) -> str:
        return re.sub(r"\s+", " ", _normalize_text(t).replace("|", " ")).strip()

    norm = norm_row(target)
    rows: list[tuple] = []
    for table in doc.tables:
        for row_index, row in enumerate(table.rows):
            joined = " | ".join(c.text.strip() for c in row.cells)
            if not joined:
                continue
            row_norm = norm_row(joined)
            if row_norm == norm:
                rows.append((len(row_norm), joined, table, row_index))
            elif norm and norm in row_norm:
                rows.append((len(row_norm), joined, table, row_index))
    if not rows:
        raise EditValidationError(
            f"未找到与「{target[:50]}」匹配的段落或表格内容，已中止本次修改。"
        )
    rows.sort(key=lambda x: x[0])
    if len(rows) == 1:
        return "table_row", (rows[0][2], rows[0][3]), "表格行匹配"
    # 多行匹配：只有最短的一行单独最短时才可以选择
    if len(rows[0][1]) < len(rows[1][1]):
        return "table_row", (rows[0][2], rows[0][3]), "表格行匹配（取最短）"
    raise AmbiguousTargetError(
        f"目标「{target[:40]}」匹配到多个表格行，请选择：",
        [r[1][:60] for r in rows[:5]],
    )




def _heading_level(paragraph: Paragraph) -> int:
    """返回段落标题级别（Heading 1..4），非标题返回 0。"""
    style = paragraph.style
    name = (style.name or "").strip()
    if name.lower().startswith("heading"):
        digits = name[7:].strip()
        if digits.isdigit():
            return int(digits)
    base = style.base_style
    while base is not None:
        bname = (base.name or "").strip()
        if bname.lower().startswith("heading"):
            digits = bname[7:].strip()
            if digits.isdigit():
                return int(digits)
        base = base.base_style
    return 0


def _first_heading_level(text: str) -> int:
    """返回文本首个非空行是否为标题；是则返回级别，否则 0。"""
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        m = HEADING_LINE_RE.match(line)
        return min(len(m.group(1)), 4) if m else 0
    return 0


def _body_level_element(doc: Document, obj) -> Any:
    """把段落/表格对象提升为 body 的直接子元素（w:p 或 w:tbl）。"""
    elem = obj._tbl if hasattr(obj, "_tbl") else obj._p
    body = doc.element.body
    while elem.getparent() is not None and elem.getparent() is not body:
        elem = elem.getparent()
    return elem


def _section_boundary(doc: Document, element, new_heading_level: int):
    """当插入内容以标题开头且层级不深于所在章节时，返回应插到其【之前】的边界元素；否则 None（紧邻插入）。"""
    containing = 0
    body = doc.element.body
    for child in body.iterchildren():
        if child is element:
            break
        if child.tag == qn("w:p"):
            lvl = _heading_level(Paragraph(child, doc))
            if lvl:
                containing = lvl
    if containing == 0 or new_heading_level > containing:
        return None
    after = False
    for child in body.iterchildren():
        if child is element:
            after = True
            continue
        if not after:
            continue
        if child.tag == qn("w:p"):
            lvl = _heading_level(Paragraph(child, doc))
            if lvl and lvl <= new_heading_level:
                return child
    return None


def _make_block_elements(doc: Document, text: str) -> list:
    """按行创建并填充段落元素，先追加到文档末尾（随后移动到目标位置）。"""
    lines = [line.rstrip() for line in text.splitlines()]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    if not lines:
        return [doc.add_paragraph()._p]
    elements = []
    for line in lines:
        new_p = doc.add_paragraph()
        _fill_line(new_p, line)
        elements.append(new_p._p)
    return elements


def _insert_after_element(doc: Document, element, text: str) -> None:
    """在 body 级元素（段落/表格）之后插入内容。

    - 新内容以标题开头且层级不深于所在章节时，自动移到该章节末尾（防止打断子节）；
    - 其余情况紧邻元素之后插入。
    """
    elements = _make_block_elements(doc, text)
    new_heading_level = _first_heading_level(text)
    if new_heading_level:
        boundary = _section_boundary(doc, element, new_heading_level)
        if boundary is not None:
            for e in elements:
                boundary.addprevious(e)
            return
    last = element
    for e in elements:
        last.addnext(e)
        last = e


def _insert_block_after_table(doc: Document, table, text: str) -> None:
    """在表格所在章节之后插入内容（保持向后兼容）。"""
    _insert_after_element(doc, _body_level_element(doc, table), text)


def _insert_block_after(doc: Document, anchor: Paragraph, text: str) -> None:
    """在段落后插入内容（新章节自动放到所在章节末尾，避免打断子节）。"""
    _insert_after_element(doc, _body_level_element(doc, anchor), text)

# ---------------------------------------------------------------- 落笔
def _preview(text: str, limit: int = 24) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    return text[:limit] + ("…" if len(text) > limit else "")


def _insert_paragraph_after(anchor: Paragraph) -> Paragraph:
    new_p = anchor._p.makeelement(qn("w:p"), {})
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    _copy_ppr(anchor._p, new_p)
    return para


def _try_set_style(paragraph: Paragraph, style: str) -> None:
    try:
        paragraph.style = style
    except KeyError:
        pass


def _clone_rpr_from_run(run) -> Any:
    """复制 run 的 rPr（字体/字号/颜色），供新 run 继承。"""
    import copy as _copy
    rpr = run._element.rPr
    return _copy.deepcopy(rpr) if rpr is not None else None


def _apply_rpr(dst_run, rpr) -> None:
    """把 rPr 元素挂到 run 上（跳过已存在的标签，避免覆盖显式加粗/斜体）。"""
    if rpr is None:
        return
    new_rpr = dst_run._element.get_or_add_rPr()
    for child in rpr:
        if new_rpr.find(child.tag) is None:
            new_rpr.append(child)


def _copy_ppr(src_p, dst_p) -> None:
    """复制段落格式（缩进/行距/对齐/边框），排除编号与大纲级别。"""
    import copy as _copy
    src_ppr = src_p.pPr
    if src_ppr is None:
        return
    dst_ppr = dst_p.get_or_add_pPr()
    for child in src_ppr:
        tag = child.tag
        if tag in (qn("w:numPr"), qn("w:outlineLvl"), qn("w:sectPr")):
            continue
        dst_ppr.append(_copy.deepcopy(child))


def _fill_line(paragraph: Paragraph, line: str, fmt_rpr=None) -> None:
    text = line
    mh = HEADING_LINE_RE.match(line)
    mb = BULLET_RE.match(line)
    mn = NUMBER_RE.match(line)
    if mh:
        text = mh.group(2)
        level = min(len(mh.group(1)), 4)
        _try_set_style(paragraph, f"Heading {level}")
    elif mb:
        text = mb.group(1)
        _try_set_style(paragraph, "List Bullet")
    elif mn:
        text = mn.group(1)
        _try_set_style(paragraph, "List Number")

    def _mk(text_: str, bold: bool = False, italic: bool = False, code: bool = False):
        run = paragraph.add_run(text_)
        _apply_rpr(run, fmt_rpr)
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        if code:
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        return run

    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            _mk(token[2:-2], bold=True)
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            _mk(token[1:-1], italic=True)
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            _mk(token[1:-1], code=True)
        else:
            _mk(token)


def _render_block_into(anchor: Paragraph, text: str, fmt_rpr=None) -> None:
    """第一行替换目标段落，其余行作为新段落插入其后。"""
    lines = [line.rstrip() for line in text.splitlines()]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    if not lines:
        anchor.add_run("")
        return
    _fill_line(anchor, lines[0], fmt_rpr=fmt_rpr)
    last = anchor
    for line in lines[1:]:
        new_p = _insert_paragraph_after(last)
        _fill_line(new_p, line, fmt_rpr=fmt_rpr)
        last = new_p




def _replace_paragraph(para: Paragraph, new_text: str, style: Optional[str]) -> None:
    # 先保存原段落第一个 run 的字体格式，替换后新内容继承（避免字体粗细/字号错乱）
    fmt_rpr = _clone_rpr_from_run(para.runs[0]) if para.runs else None
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    if style:
        if style not in VALID_STYLES:
            raise EditValidationError(f"不支持的样式：{style}")
        _try_set_style(para, style)
    _render_block_into(para, new_text, fmt_rpr=fmt_rpr)
 
 
def _reorder_structure_first(resolved):
    """结构操作优先于内容操作（officecli：先结构后内容）。

    结构操作：set_style / delete / 新内容以标题开头（replace / insert_after）。
    仅在所有操作目标互不冲突（没有 delete 与其他操作目标相互包含）时重排；
    一旦存在冲突就保持原顺序，避免锚点失效，保证安全。
    """
    norm_targets = [_normalize_text(str(op.get("target", ""))) for op, *_ in resolved]
    for i, (op, *_) in enumerate(resolved):
        if op["op"] != "delete":
            continue
        t_i = norm_targets[i]
        if not t_i:
            continue
        for j, t_j in enumerate(norm_targets):
            if i == j or not t_j:
                continue
            if t_i in t_j or t_j in t_i:
                return resolved  # 存在交叉目标，保持原序

    def _is_structure(item) -> bool:
        op = item[0]
        if op["op"] in ("set_style", "delete"):
            return True
        return bool(_first_heading_level(op.get("new_text", "")))

    return sorted(resolved, key=lambda it: 0 if _is_structure(it) else 1)


def apply_operations(
    doc: Document,
     operations: list[dict],
    log: Callable[[str], None] = print,
    resolve_ambiguous: Optional[Callable[[dict, list[str]], Optional[str]]] = None,
) -> list[str]:
    """先全量校验目标，全部通过才逐条落笔（保证不写一半）。

    resolve_ambiguous(op, candidates) 在目标有歧义时调用，返回用户选中的候选文本，
    返回 None 表示取消（整体中止，不写盘）。
    """
    resolved: list[tuple[dict, str, Any, str]] = []
    for op in operations:
        try:
            kind, obj, how = _find_target(doc, str(op["target"]))
        except AmbiguousTargetError as exc:
            if resolve_ambiguous is None:
                raise
            chosen = resolve_ambiguous(op, exc.candidates)
            if not chosen:
                raise EditValidationError(f"目标「{op['target'][:40]}」有歧义且未选择，已中止本次修改。")
            kind, obj, how = _find_target(doc, chosen)
        resolved.append((op, kind, obj, how))
 
    # 结构优先（officecli：先结构后内容）：删除/改样式/新增章节等结构操作先落笔，
    # 纯文本替换与插入后落笔；仅在目标互不冲突时重排，冲突则保持原序（最安全）。
    resolved = _reorder_structure_first(resolved)

    applied: list[str] = []
    for op, kind, obj, how in resolved:
        op_type = op["op"]
        target = str(op["target"])
        new_text = op.get("new_text", "")
        style = op.get("style")
        if kind == "table_row":
            table, row_index = obj
            row = table.rows[row_index]
            if op_type == "insert_after":
                _insert_block_after_table(doc, table, new_text)
                applied.append(f"在表格「{_preview(target)}」后插入内容（{how}）")
            elif op_type == "delete":
                row._tr.getparent().remove(row._tr)
                applied.append(f"删除表格行「{_preview(target)}」（{how}）")
            elif op_type == "replace":
                lines = [ln.rstrip() for ln in (new_text or "").splitlines() if ln.strip()]
                if not lines:
                    raise EditValidationError("替换表格行需要提供新内容。")
                cells = row.cells
                for cell in cells:
                    for para in list(cell.paragraphs):
                        for run in list(para.runs):
                            run._element.getparent().remove(run._element)
                fmt_rpr = None
                if cells and cells[0].paragraphs and cells[0].paragraphs[0].runs:
                    fmt_rpr = _clone_rpr_from_run(cells[0].paragraphs[0].runs[0])
                for c_i, cell in enumerate(cells):
                    p0 = cell.paragraphs[0]
                    if c_i < len(lines):
                        _fill_line(p0, lines[c_i], fmt_rpr=fmt_rpr)
                    else:
                        p0.add_run("")
                applied.append(f"替换表格行「{_preview(target)}」（{how}）")
            else:
                raise EditValidationError(
                    f"表格行目标不支持 {op_type}；替换请用 replace，删除请用 delete。"
                )
            continue
        para: Paragraph = obj
        if op_type == "replace":
            _replace_paragraph(para, new_text, style)
            applied.append(f"替换「{_preview(target)}」（{how}）")
        elif op_type == "insert_after":
            _insert_block_after(doc, para, new_text)
            applied.append(f"在「{_preview(target)}」后插入内容（{how}）")
        elif op_type == "delete":
            para._p.getparent().remove(para._p)
            applied.append(f"删除「{_preview(target)}」（{how}）")
        elif op_type == "set_style":
            if style not in VALID_STYLES:
                raise EditValidationError(f"不支持的样式：{style}")
            _try_set_style(para, style)
            applied.append(f"「{_preview(target)}」样式改为 {style}（{how}）")
        else:
            raise EditValidationError(f"未知操作类型：{op_type}")
    return applied


# ---------------------------------------------------------------- 程序化结构校验（不依赖 LLM）
def _check_structural(state: dict[str, Any]) -> list[str]:
    """校验：原标题完整性 + 替换/删除目标是否真的生效。返回问题列表（空=无问题）。"""
    doc: Document = state["doc"]
    ops: list[dict] = state["plan"]["operations"]
    headings_before = state.get("headings_before", [])
    issues: list[str] = []

    markdown, _ = docx_to_markdown(doc)
    headings_after = set(line for line in markdown.splitlines() if line.startswith("#"))

    # 被操作指向过的标题（允许消失）
    targeted: set[str] = set()
    for op in ops:
        norm = _normalize_text(str(op.get("target", "")))
        if not norm:
            continue
        for heading in headings_before:
            if _normalize_text(heading) == norm or norm in _normalize_text(heading):
                targeted.add(heading)

    for heading in headings_before:
        if heading not in headings_after and heading not in targeted:
            issues.append(f"原有标题「{heading[:40]}」在修改后消失，且没有操作明确指向它")

    # 替换/删除的目标不应再作为整段存在
    for op in ops:
        if op["op"] not in ("replace", "delete"):
            continue
        norm = _normalize_text(str(op.get("target", "")))
        if not norm:
            continue
        still = [p for p in _iter_all_paragraphs(doc) if _normalize_text(p.text) == norm]
        if still:
            issues.append(f"操作[{op['op']}]的目标「{str(op['target'])[:30]}」修改后仍整段存在，可能未生效")

    return issues


# ---------------------------------------------------------------- AI 复核（建议模式）
def _verification_review(
    state: dict[str, Any],
    llm: LLMClient,
    log: Callable[[str], None],
    max_rounds: int = 2,
    auto_fix: bool = False,
) -> bool:
    """让 AI 对照要求复核（完整文档 + 已执行操作清单）。

    默认建议模式（只报告不落笔）；auto_fix=True 时应用建议并再次复核。
    """
    doc: Document = state["doc"]
    instruction: str = state["instruction"]
    headings_before = state.get("headings_before", [])
    ops_done = state["plan"]["operations"]
    ops_summary = "\n".join(
        f"- [{op.get('op')}] 目标「{str(op.get('target', ''))[:60]}」"
        + (f" → {str(op.get('new_text', ''))[:80]}" if op.get("new_text") else "")
        for op in ops_done
    )
    for _round in range(1, max_rounds + 1):
        markdown, _ = docx_to_markdown(doc)
        headings_after = [line for line in markdown.splitlines() if line.startswith("#")]
        # 只提取改动目标附近的局部上下文（前后各 2 行），大幅省 token
        lines = markdown.splitlines()
        region = set()
        for op in ops_done:
            target = _normalize_text(str(op.get("target", "")))
            if not target:
                continue
            for i, line in enumerate(lines):
                if target in _normalize_text(line):
                    for j in range(max(0, i - 2), min(len(lines), i + 3)):
                        region.add(j)
        context_lines = [lines[j] for j in sorted(region)]
        context = "\n".join(context_lines)
        if len(context) > 6000:
            context = context[:6000] + "\n\n……（改动区域上下文较长，已截断）……"
        if not context:
            context = "（未能定位改动段落，请结合标题清单判断修改是否生效）"
        messages = [
            {"role": "system", "content": CRITIC_PROMPT},
            {
                "role": "user",
                "content": (
                    f"【用户修改要求】\n{instruction}\n\n"
                    f"【已执行的操作】\n{ops_summary or '（无）'}\n\n"
                    f"【修改前标题】\n" + ("\n".join(headings_before) or "（无）") + "\n\n"
                    f"【修改后标题】\n" + ("\n".join(headings_after) or "（无）") + "\n\n"
                    f"【改动区域上下文】\n{context}"
                ),
            },
        ]
        try:
            data = llm.chat_json(messages, temperature=0.1, max_tokens=8192)
        except (LLMError, Exception) as exc:  # noqa: BLE001
            log(f"   ⚠ AI 复核调用失败（{exc}），跳过复核。")
            return []
        if data.get("satisfied"):
            log("   ✓ AI 复核通过：修改满足要求，原有标题完整。")
            return True
        reason = str(data.get("reason", "")).strip()
        issues = data.get("issues") or []
        log(f"   ↻ AI 复核意见：{reason or '存在待确认问题'}")
        for issue in issues[:5]:
            log(f"      · {str(issue)[:120]}")
        suggested = _normalize_ops(data.get("operations"))
        if not suggested:
            log("   复核未给出补充操作，以当前内容为准。")
            return True
        log(f"   复核建议 {len(suggested)} 项补充操作" + ("（将自动应用）" if auto_fix else "（需人工确认，未自动执行）") + ":")
        for op in suggested:
            extra = op.get("new_text", op.get("style", ""))
            log(f"      - [{op['op']}] 目标「{str(op['target'])[:40]}」"
                + (f" → {str(extra)[:60]}" if extra else ""))
        if not auto_fix:
            return False
        try:
            fixed = apply_operations(doc, suggested, log)
            for item in fixed:
                log(f"      ✓ {item}")
        except EditValidationError as exc:
            log(f"   ⚠ 复核建议未通过校验（{exc}），保持现有修改。")
            return True
    log("   复核轮次已达上限，以当前内容为准。")
    return True


# ---------------------------------------------------------------- 编排（两阶段：prepare / finalize）
def prepare_edit(
    src_path: Path,
    instruction: str,
    config: Config,
    memory: Memory,
    llm: Optional[LLMClient] = None,
    log: Callable[[str], None] = print,
    reference_files: Optional[list] = None,
) -> dict[str, Any]:
    """第一阶段：读取文档、解析指令、生成修改计划。返回中间状态。

    reference_files：可选的参考文件（图片/文本/csv/docx），内容会注入计划提示，
    供 AI 结合参考资料修改（例如按参考文档的格式/数据改当前文档）。
    """
    src = Path(src_path)
    if not src.exists():
        raise FileNotFoundError(f"文件不存在：{src}")
    if src.suffix.lower() != ".docx":
        raise ValueError("仅支持编辑 .docx 文件。")

    memory.add_user(instruction)
    llm = llm or LLMClient(config)
    log(f"① 读取文档：{src.name}")
    doc = Document(str(src))
    markdown, info = docx_to_markdown(doc)
    log(f"   已提取 {len(info)} 个段落/单元格，共 {len(markdown)} 字符")

    # 文档内要求：批注 / 页眉页脚 / 文本框（用户的要求常写在文档里，必须让 AI 读到）
    doc_extra = ""
    try:
        from .extractor import doc_requirements as _doc_req
        doc_extra = _doc_req(src)
        if doc_extra:
            log(f"   已读取文档内要求（批注/页眉页脚/文本框）：{len(doc_extra)} 字符")
    except Exception:
        pass

    # 参考文件注入（编辑模式也支持附参考资料/模板）
    refs_extra = ""
    try:
        from .refdocs import collect_reference_context as _collect_refs
        _refs = _collect_refs(reference_files, log=log)
        if _refs.get("names"):
            refs_extra = (
                "\n\n【参考文件内容（修改时须结合参考，作为格式/数据/内容依据）】\n"
                + (_refs.get("writer_context") or "")[:5000]
            )
    except Exception:
        pass

    log("② 解析编辑指令，生成修改计划 ...")
    plan = plan_edit(instruction, markdown, memory.context_text(), llm,
                     doc_requirements=doc_extra + refs_extra)
    log(f"   计划说明：{plan.get('summary', '')}（{len(plan['operations'])} 项操作）")
    for op in plan["operations"]:
        extra = op.get("new_text", op.get("style", ""))
        log(f"   - [{op['op']}] 目标「{_preview(op['target'], 40)}」 {('→ ' + _preview(str(extra), 30)) if extra else ''}")
    headings_before = [line for line in markdown.splitlines() if line.startswith("#")]
    return {"src": src, "doc": doc, "plan": plan, "instruction": instruction, "llm": llm,
            "headings_before": headings_before}


def finalize_edit(
    state: dict[str, Any],
    config: Config,
    memory: Memory,
    output_dir: Optional[Path] = None,
    log: Callable[[str], None] = print,
    save_as_new: bool = True,
    verify: bool = True,
    auto_fix: bool = False,
    resolve_ambiguous: Optional[Callable[[dict, list[str]], Optional[str]]] = None,
    on_warnings: Optional[Callable[[list[str]], bool]] = None,
    version_keep: bool = False,
) -> Path:
    """第二阶段：应用修改 -> 结构校验 -> AI 复核（建议） -> 备份 -> 保存 -> 完整性校验。"""
    src: Path = state["src"]
    doc: Document = state["doc"]
    plan: dict = state["plan"]
    llm: LLMClient = state["llm"]
    operations = plan["operations"]

    log("③ 校验并应用修改（精确→包含→近似三级匹配，不匹配即中止）...")
    applied = apply_operations(doc, operations, log, resolve_ambiguous=resolve_ambiguous)
    for item in applied:
        log(f"   ✓ {item}")

    # 程序化结构校验（不依赖 LLM）
    issues = _check_structural(state)
    if issues:
        log("   ⚠ 结构校验发现以下问题：")
        for issue in issues:
            log(f"      · {issue}")
        if on_warnings is not None and not on_warnings(issues):
            raise EditValidationError("检测到结构风险，已中止保存（原文件未受影响）。")
    else:
        log("   ✓ 结构校验通过：原标题完整，操作均已生效。")

    if verify:
        log("④ AI 复核：检查修改是否真正满足要求 ...")
        _verification_review(state, llm, log, auto_fix=auto_fix)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = Path(output_dir or config.output_dir)
    if version_keep:
        # 双版本模式：始终只保留「原始版 + 完成版」两个版本，方便回退
        ver_dir = out_dir / "versions"
        ver_dir.mkdir(parents=True, exist_ok=True)
        base = src.stem
        for suf in ("_完成版", "_原始版", "_修改版"):
            if base.endswith(suf):
                base = base[: -len(suf)]
                break
        orig_v = ver_dir / f"{base}_原始版.docx"
        done_v = ver_dir / f"{base}_完成版.docx"
        shutil.copy2(src, orig_v)
        log(f"   已保存原始版：{orig_v.name}")
        for old_file in list(ver_dir.glob(f"{base}_*")):
            if old_file.name not in (orig_v.name, done_v.name):
                try:
                    old_file.unlink()
                    log(f"   已移除最早版本：{old_file.name}")
                except OSError:
                    pass
        out_path = done_v
        backup_path = orig_v
    else:
        backup_dir = out_dir / "backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        backup_path = backup_dir / f"{src.stem}_备份_{stamp}{src.suffix}"
        shutil.copy2(src, backup_path)
        log(f"   已备份原文件：{backup_path}")

        if save_as_new:
            out_path = src.parent / f"{src.stem}_修改版_{stamp}{src.suffix}"
        else:
            out_path = src
    doc.save(str(out_path))
    log(f"   已保存：{out_path}")
 
    # 程序化质量门禁（officecli Verification Gate）：自动修复安全项并复检，
    # 覆盖标题层级/空段污染/Markdown 残留/字体/表格边框/内容完整度，不消耗 token
    try:
        from .verify import fix_document as _fix_doc
        log("⑤ 质量检查并自动修复 ...")
        _fix_doc(out_path, log=log)
    except Exception as _exc:  # noqa: BLE001
        log(f"   ⚠ 质量检查跳过（{_exc}）")

    # 完整性校验
    check = Document(str(out_path))
    if len(check.paragraphs) == 0:
        raise EditValidationError("保存后的文档为空，已中止（原文件未受影响）。")
    log("   ✓ 完整性校验通过")

    memory.add_result(
        {
            "action": "edit",
            "title": src.name,
            "file": str(out_path),
            "backup": str(backup_path),
            "original": str(orig_v) if version_keep else "",
            "summary": plan.get("summary", ""),
            "operations": len(operations),
            "verified": verify,
        }
    )
    log(f"✔ 编辑完成：{out_path}")
    return out_path


def edit_document(
    src_path: Path,
    instruction: str,
    config: Config,
    memory: Memory,
    llm: Optional[LLMClient] = None,
    output_dir: Optional[Path] = None,
    log: Callable[[str], None] = print,
    confirm: Optional[Callable[[dict], bool]] = None,
    save_as_new: bool = True,
    dry_run: bool = False,
    verify: bool = True,
    auto_fix: bool = False,
    resolve_ambiguous: Optional[Callable[[dict, list[str]], Optional[str]]] = None,
    on_warnings: Optional[Callable[[list[str]], bool]] = None,
    reference_files: Optional[list] = None,
) -> Optional[Path]:
    """完整编辑流程（prepare + 可选确认 + finalize）。dry_run 只预览不落笔。"""
    state = prepare_edit(src_path, instruction, config, memory, llm=llm, log=log,
                         reference_files=reference_files)
    if dry_run:
        log("   [dry-run] 仅预览修改计划，未做任何修改。")
        return None
    if confirm is not None and not confirm(state["plan"]):
        log("   已取消，未做任何修改。")
        return None
    return finalize_edit(
        state, config, memory, output_dir=output_dir, log=log,
        save_as_new=save_as_new, verify=verify, auto_fix=auto_fix,
        resolve_ambiguous=resolve_ambiguous, on_warnings=on_warnings,
    )
