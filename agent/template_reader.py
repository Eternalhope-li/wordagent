from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

_EMU_PER_CM = 360000


def _cm(emu) -> float:
    try:
        return round(float(emu) / _EMU_PER_CM, 2)
    except (TypeError, ValueError):
        return 0.0


def _style_fonts(style):
    """返回样式 (ascii/latin 字体, eastAsia 中文字体)；主题字体或缺失返回 None。"""
    if style is None:
        return None, None
    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        return None, None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return None, None
    ascii_font = rfonts.get(qn("w:ascii")) or rfonts.get(qn("w:hAnsi"))
    ea_font = rfonts.get(qn("w:eastAsia"))
    return ascii_font, ea_font


def extract_template(template_path: str | Path) -> dict:
    """从模板 docx 提取结构信息，返回 dict（各字段均可缺失，调用方需容错）。"""
    doc = Document(str(template_path))
    info: dict = {}

    # ---- 页面设置 ----
    try:
        sec = doc.sections[0]
        page = {
            "width_cm": _cm(sec.page_width) or 21.0,
            "height_cm": _cm(sec.page_height) or 29.7,
            "top_cm": _cm(sec.top_margin) or 2.5,
            "bottom_cm": _cm(sec.bottom_margin) or 2.5,
            "left_cm": _cm(sec.left_margin) or 2.8,
            "right_cm": _cm(sec.right_margin) or 2.8,
        }
        info["page"] = page
    except Exception:
        pass

    # ---- 正文字体（优先 eastAsia 中文字体；西文字体单独记录） ----
    try:
        nf = doc.styles["Normal"]
        ascii_font, ea_font = _style_fonts(nf)
        body_font = ea_font or nf.font.name or "微软雅黑"
        info["body_font"] = body_font
        info["body_size"] = round(float(nf.font.size.pt), 1) if nf.font.size else 12.0
        if ascii_font and ascii_font != body_font:
            info["body_font_latin"] = ascii_font
    except Exception:
        pass

    # ---- 标题字体 / 颜色 / 字号 ----
    try:
        h1 = doc.styles["Heading 1"]
        _, h1_ea = _style_fonts(h1)
        info["heading_font"] = h1_ea or h1.font.name or "微软雅黑"
    except Exception:
        pass
    try:
        c = doc.styles["Heading 1"].font.color
        if c is not None and c.rgb is not None:
            info["heading_color"] = str(c.rgb)
    except Exception:
        pass
    heading_sizes: dict = {}
    for lv in (1, 2, 3, 4):
        try:
            st = doc.styles[f"Heading {lv}"].font
            if st.size is not None:
                heading_sizes[lv] = round(float(st.size.pt), 1)
        except Exception:
            pass
    if heading_sizes:
        info["heading_sizes"] = heading_sizes

    # ---- 章节结构（Heading 样式 + 加粗/居中/序号开头的普通段落标题） ----
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _NUM_RE = re.compile(r"^[（(]?[0-9一二三四五六七八九十]+[、．.）)]")
    sections: list[dict] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text or len(text) > 40:
            continue
        name = (p.style.name or "") if p.style is not None else ""
        m = re.match(r"(?i)^\s*heading\s*(\d)\s*$", name.strip())
        if m:
            sections.append({"heading": text, "level": int(m.group(1))})
            continue
        # 普通段落形似标题：加粗 / 居中 / 序号开头（排除正文句子）
        try:
            bold = any(r.font.bold for r in p.runs)
            centered = p.alignment == WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            bold, centered = False, False
        if (bold or centered or _NUM_RE.match(text)) and not text.rstrip().endswith(("。", "；", "：", "，", "！", "？")):
            sections.append({"heading": text, "level": 1})
    if sections:
        # 过滤文档大标题（首个居中/大字号段落，如「实 验 报 告」），避免当成章节
        _TITLE_KW = ("报告", "说明书", "设计", "方案", "论文", "作业", "总结")
        filtered = []
        for idx, s in enumerate(sections):
            t = s["heading"]
            is_doc_title = (
                idx == 0
                and len(t) <= 16
                and any(k in t for k in _TITLE_KW)
            )
            if not is_doc_title:
                filtered.append(s)
        info["sections"] = filtered

    return info


def structure_text(info: dict) -> str:
    """把模板章节结构格式化为提示词文本。"""
    secs = info.get("sections") or []
    if not secs:
        return ""
    lines = []
    for s in secs:
        lines.append(f"{'  ' * (s['level'] - 1)}- [{s['level']}] {s['heading']}")
    return "\n".join(lines)


def to_renderer_preset(info: dict) -> dict:
    """把模板信息转成 renderer 可用的 preset（只含有效字段）。"""
    preset: dict = {}
    for key in ("body_font", "body_font_latin", "heading_font", "body_size",
                "heading_color", "heading_sizes", "page"):
        if info.get(key):
            preset[key] = info[key]
    return preset

def full_text(template_path: str | Path) -> str:
    """提取模板全部非空文本（正文段落 + 表格单元格 + 页眉页脚），供 AI 读取模板内要求。"""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    doc = Document(str(template_path))
    lines: list[str] = []
    body = doc.element.body

    def walk(el) -> None:
        for child in el.iterchildren():
            if child.tag == qn("w:p"):
                t = Paragraph(child, doc).text.strip()
                if t:
                    lines.append(t)
            elif child.tag == qn("w:tbl"):
                tbl = Table(child, doc)
                for row in tbl.rows:
                    cells = [c.text.replace("\n", " ").strip() for c in row.cells]
                    if any(cells):
                        lines.append(" | ".join(cells))
                walk(child)

    walk(body)
    for sec in doc.sections:
        for p_ in sec.header.paragraphs:
            t = p_.text.strip()
            if t:
                lines.append("（页眉）" + t)
        for p_ in sec.footer.paragraphs:
            t = p_.text.strip()
            if t:
                lines.append("（页脚）" + t)
    return "\n".join(lines)
