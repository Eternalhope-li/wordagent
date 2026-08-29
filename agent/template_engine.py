"""模板驱动编辑引擎：理解模板 -> AI 填内容 -> 格式 100% 继承模板。

设计原则（参考 GitHub 成熟方案 docxtpl / report-forge / DocuFiller）：
- 模板是母版：格式（页面/样式/字体/段落/表格/页眉页脚）一律保留，绝不重建
- 先理解模板：解析每个段落/表格的文本、样式、是否标题、是否空白待填
- AI 只输出「哪里填什么」，不输出任何排版/格式指令
- 落笔复制相邻段落格式（deepcopy rPr/pPr），表格新行复制模板行格式
"""
from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path
from typing import Any, Callable, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# ---------- 模板解析 ----------

_EMPTY_PATTERNS = (
    re.compile(r"^\s*[\[\{（(]?\s*(?:填写|内容|待填|此处|说明|备注|正文)?\s*[\]\}）)]?\s*$"),
    re.compile(r"^[＿_—\-]{3,}$"),
    re.compile(r"^【[^】]{0,12}】"),
    re.compile(r"[＿_]{3,}"),
)

_PLACEHOLDER_RE = re.compile(r"[\[\{（(]\s*([^\]\}）)]{1,30}?)\s*[\]\}）)]")

_HINT_RE = re.compile(r"^[（(][^）)]{2,80}[）)]$")

_LABEL_WORDS = (
    "姓名", "学号", "班级", "院系", "专业", "课程", "实验项目", "项目名称", "预习日期",
    "日期", "实验学时", "学时", "实验性质", "性质", "教师签名", "成绩", "批阅日期",
    "签名", "签字", "实验目的", "目的及要求", "实验设备", "实验环境", "实验内容",
    "实验步骤", "实验原理", "实验方法", "预习任务", "总结", "批阅意见", "心得体会",
    "问题", "环境", "要求", "实验名称", "班级学号",
)

_INFO_WORDS = (
    "姓名", "学号", "班级", "院系", "专业", "签名", "签字", "批阅", "成绩",
    "时间", "教师", "指导老师",
)

def _para_style(para: Paragraph) -> str:
    try:
        return (para.style.name if para.style is not None else "Normal") or "Normal"
    except Exception:
        return "Normal"

def _para_text(para: Paragraph) -> str:
    return "".join(run.text for run in para.runs)

def _is_empty_para(para: Paragraph) -> bool:
    text = _para_text(para).strip()
    if not text:
        return True
    if _HINT_RE.match(text):
        return True
    return any(p.match(text) for p in _EMPTY_PATTERNS)

def _heading_level(para: Paragraph) -> int:
    name = _para_style(para).lower()
    for part in name.split():
        if part.isdigit():
            return int(part)
    if name.startswith(("heading", "标题")):
        return 1
    ppr = para._p.pPr
    if ppr is not None:
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is not None:
            try:
                return int(outline.get(qn("w:val"), "0")) + 1
            except ValueError:
                return 1
    return 0

def _looks_like_title(para: Paragraph) -> bool:
    """普通段落形似标题：加粗 / 居中 / 序号开头 / 短行。"""
    text = _para_text(para).strip()
    if not text or len(text) > 40:
        return False
    try:
        bold = any(r.font.bold for r in para.runs)
        centered = para.alignment is not None and str(para.alignment) == "CENTER (1)"
    except Exception:
        bold, centered = False, False
    if bold or centered:
        return True
    return bool(re.match(r"^[（(]?[0-9一二三四五六七八九十]+[、．.）)]", text))

def _is_label_text(text: str) -> bool:
    """判断一段文本是否为表单标签（栏目名）。"""
    t = re.sub(r"\s+", "", text or "").strip()
    if not t or len(t) > 30:
        return False
    if re.search(r"[，。；！？]", t):
        return False
    if t.endswith(("：", ":")):
        return True
    return any(w in t for w in _LABEL_WORDS) and len(t) <= 12


def _is_info_label(label: str) -> bool:
    """信息栏标签（签名/日期/成绩等，默认不主动填写）。"""
    return any(w in re.sub(r"\s+", "", label or "") for w in _INFO_WORDS)


def _split_label_hint(text: str):
    """拆开「总结与实验体会（对实验结果进行分析…）」这类标签+提示。"""
    m = re.match(r"^(.{2,24}?)[（(]([^）)]{0,60})[）)]$", (text or "").strip())
    if m:
        label, hint = m.group(1), m.group(2)
        if _is_label_text(label) or any(w in label for w in _LABEL_WORDS):
            return label.rstrip("：:"), hint
    return None


def _load_comments(template_path) -> dict:
    """读取 docx 批注：{comment_id: 批注文本}。"""
    import zipfile
    cxml = ""
    try:
        with zipfile.ZipFile(str(template_path)) as zf:
            if "word/comments.xml" in zf.namelist():
                cxml = zf.read("word/comments.xml").decode("utf-8", "ignore")
    except (KeyError, zipfile.BadZipFile, OSError):
        pass
    comments: dict[str, str] = {}
    if cxml:
        try:
            from lxml import etree as _et
            croot = _et.fromstring(cxml.encode("utf-8"))
            for c in croot.iter():
                if c.tag == qn("w:comment"):
                    cid = c.get(qn("w:id"))
                    txt = "".join(t.text or "" for t in c.iter(qn("w:t")))
                    txt = re.sub(r"\s+", " ", txt).strip()
                    if cid and txt:
                        comments[cid] = txt
        except Exception:
            pass
    return comments


def _para_comment(para, comments_by_id: dict) -> str:
    """段落批注文本（直接检查段落内 commentRangeStart，无身份匹配问题）。"""
    try:
        ids = [c.get(qn("w:id")) for c in para._p.iter(qn("w:commentRangeStart"))]
    except Exception:
        return ""
    return "；".join(comments_by_id.get(i, "") for i in ids if comments_by_id.get(i))


def _table_is_form(table) -> bool:
    """表单型表格：有合并单元格，或首行是标签行。"""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is not None and (
                tcPr.find(qn("w:gridSpan")) is not None
                or tcPr.find(qn("w:vMerge")) is not None
            ):
                return True
    if table.rows:
        first = [c.text.strip() for c in table.rows[0].cells]
        non_empty = [t for t in first if t]
        if non_empty:
            labely = sum(1 for t in non_empty if _is_label_text(t))
            if labely >= max(1, len(non_empty) * 0.5):
                return True
    return False


def _cell_plain(cell) -> str:
    """单元格纯文本（合并单元格按一份算）。"""
    return "".join(p.text for p in cell.paragraphs).strip()


def _extract_form_slots(table, comments_by_id: dict) -> list[dict]:
    """识别表单表格内的填写位（cell_paras / cell_value）。"""
    slots: list[dict] = []
    for r_i, row in enumerate(table.rows):
        seen = set()
        for c_i, cell in enumerate(row.cells):
            if id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))
            paras = cell.paragraphs
            texts = [_para_text(p).strip() for p in paras]
            for p_i, t in enumerate(texts):
                if not t:
                    continue
                label, hint = None, None
                if _is_label_text(t):
                    label = t.rstrip("：:")
                else:
                    sp = _split_label_hint(t)
                    if sp:
                        label, hint = sp
                if label is None:
                    continue
                blank = 0
                k = p_i + 1
                while k < len(texts) and (not texts[k] or _HINT_RE.match(texts[k])):
                    blank += 1
                    k += 1
                # 标签后有空/提示段，或紧邻另一个标签（内容插到标签后）→ 填写位
                if blank > 0 or (k < len(texts) and _is_label_text(texts[k])):
                    slots.append({
                        "kind": "cell_paras", "row": r_i, "col": c_i,
                        "para_index": p_i, "label": label, "hint": hint or "",
                        "blank_after": blank, "info_field": _is_info_label(label),
                        "comment": _para_comment(paras[p_i], comments_by_id),
                    })
                    continue
                if c_i + 1 < len(row.cells):
                    right = row.cells[c_i + 1]
                    if id(right._tc) != id(cell._tc) and not _cell_plain(right):
                        slots.append({
                            "kind": "cell_value", "row": r_i, "col": c_i,
                            "value_row": r_i, "value_col": c_i + 1,
                            "label": label, "hint": hint or "",
                            "info_field": _is_info_label(label),
                            "comment": _para_comment(paras[p_i], comments_by_id),
                        })
    return slots


def _extract_para_slots(blocks: list[dict]) -> list[dict]:
    """正文里的「标签段：+ 空段」填写位。"""
    slots: list[dict] = []
    for i, b in enumerate(blocks):
        if b["kind"] != "paragraph":
            continue
        t = (b["text"] or "").strip()
        if not t or len(t) > 30:
            continue
        label, hint = None, None
        if _is_label_text(t):
            label = t.rstrip("：:")
        else:
            sp = _split_label_hint(t)
            if sp:
                label, hint = sp
        if label is None:
            continue
        nxt = None
        for j in range(i + 1, len(blocks)):
            if blocks[j]["kind"] == "paragraph":
                nxt = blocks[j]
                break
        if nxt is not None and nxt.get("is_empty"):
            slots.append({
                "kind": "para_paras", "label": label, "hint": hint or "",
                "info_field": _is_info_label(label),
                "comment": b.get("comment", ""),
            })
    return slots


def _find_fill_slots(info: dict, label: str) -> list[dict]:
    """按 label 找填写位：精确同名全部返回；无同名时取最相似的一个。"""
    norm = re.sub(r"\s+", "", label).rstrip("：:")
    exact: list[dict] = []
    best, best_score = None, 0
    for t in info.get("tables", []):
        for s in t.get("slots", []):
            sn = re.sub(r"\s+", "", s["label"])
            if sn == norm:
                exact.append({**s, "table": t["table"], "slot_key": id(s)})
            elif norm and (norm in sn or sn in norm):
                score = min(len(norm), len(sn))
                if score > best_score:
                    best_score, best = score, {**s, "table": t["table"], "slot_key": id(s)}
    for s in info.get("para_slots", []):
        sn = re.sub(r"\s+", "", s["label"])
        if sn == norm:
            exact.append({**s, "slot_key": id(s)})
        elif norm and (norm in sn or sn in norm):
            score = min(len(norm), len(sn))
            if score > best_score:
                best_score, best = score, {**s, "slot_key": id(s)}
    if exact:
        return exact
    return [best] if best_score >= 2 else []


def _clone_rpr(source_rpr, drop_bold: bool = False):
    """深拷贝 rPr，可选去掉加粗。"""
    if source_rpr is None:
        return None
    new = deepcopy(source_rpr)
    if drop_bold:
        for tag in ("w:b", "w:bCs"):
            node = new.find(qn(tag))
            if node is not None:
                new.remove(node)
    return new



# ---------- 内容清理（AI 内容里的 Markdown 残留 -> 纯文本） ----------

def _clean_fill_text(text: str) -> str:
    """去掉 **加粗** / *斜体* / 代码 标记，只保留文字本身。"""
    t = str(text)
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)
    t = re.sub(r"\*([^*\n]+?)\*", r"\1", t)
    t = re.sub(r"`([^`]*)`", r"\1", t)
    return t.strip()


def _convert_list_marker(line: str) -> str:
    """把 markdown 列表前缀转成 docx 可读形式：- x -> • x，1. x -> 1. x。"""
    s = line.strip()
    m = re.match(r"^\s*[-*]\s+(.*)$", s)
    if m:
        return "• " + m.group(1).strip()
    m = re.match(r"^\s*(\d+)[.、)]\s+(.*)$", s)
    if m:
        return f"{m.group(1)}. {m.group(2).strip()}"
    return line


def _clean_fill_line(line: str) -> str:
    """单行内容：先转列表前缀，再去行内标记。"""
    return _clean_fill_text(_convert_list_marker(line))


def _insert_para_after(doc: Document, anchor_para, text: str, fmt_source, drop_bold: bool = False) -> Paragraph:
    """在锚点段后插入新段：复制格式样本的 pPr/rPr。"""
    new_p = doc.add_paragraph()
    _copy_para_format(fmt_source, new_p)
    run = new_p.add_run(text)
    src_run = fmt_source.runs[0] if fmt_source.runs else None
    if src_run is not None and src_run._element.rPr is not None:
        rpr = _clone_rpr(src_run._element.rPr, drop_bold=drop_bold)
        if rpr is not None:
            new_rpr = run._element.get_or_add_rPr()
            for child in rpr:
                if new_rpr.find(child.tag) is None:
                    new_rpr.append(child)
    anchor_para._p.addnext(new_p._p)
    return new_p


def _fill_slot(doc: Document, slot: dict, value: str) -> int:
    """把一个填写位的值写入（继承单元格/段落格式）。"""
    lines = [_clean_fill_line(ln) for ln in value.splitlines()]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    if not lines:
        return 0
    if slot["kind"] in ("cell_value",):
        table = slot["table"]
        cell = table.rows[slot["value_row"]].cells[slot["value_col"]]
        p0 = cell.paragraphs[0]
        if not _para_text(p0).strip():
            _replace_para_text_keep_format(p0, lines[0])
            anchor, fmt = p0, p0
            rest = lines[1:]
        else:
            anchor = fmt = cell.paragraphs[-1]
            rest = lines
        for line in rest:
            anchor = _insert_para_after(doc, anchor, line, fmt)
        return len(lines)
    if slot["kind"] in ("cell_paras", "para_paras"):
        if slot["kind"] == "cell_paras":
            table = slot["table"]
            cell = table.rows[slot["row"]].cells[slot["col"]]
            paras = cell.paragraphs
        else:
            paras = slot["paras"]
        # 按文本重新定位标签段（前面填写会改变段落索引）
        want = re.sub(r"\s+", "", slot["label"])
        label_p = None
        for p in paras:
            pt = re.sub(r"\s+", "", _para_text(p))
            if pt.rstrip("：:") == want or pt.startswith(want + "（") or pt.startswith(want + "("):
                label_p = p
                break
        if label_p is None:
            return 0
        # 标签后的第一个段落：空段/提示段 → 替换；紧邻标签 → 插入；已有正文 → 不动
        target, after = None, False
        for p in paras:
            if p is label_p:
                after = True
                continue
            if not after:
                continue
            pt = _para_text(p).strip()
            if not pt or _HINT_RE.match(pt):
                target = p
                break
            if _is_label_text(pt):
                break
            break  # 非空非标签正文：该位已有内容，不覆盖
        if target is not None:
            _replace_para_text_keep_format(target, lines[0])
            anchor, fmt, drop_bold = target, target, False
            rest = lines[1:]
        else:
            anchor, fmt, drop_bold = label_p, label_p, True
            rest = lines
        for line in rest:
            anchor = _insert_para_after(doc, anchor, line, fmt, drop_bold=drop_bold)
        return len(lines)
    return 0


def _copy_para_format(template_para: Paragraph, new_para: Paragraph) -> None:
    """复制段落格式（pPr）到新段落。"""
    src_pPr = template_para._p.pPr
    if src_pPr is not None:
        dst_pPr = new_para._p.get_or_add_pPr()
        for child in src_pPr:
            if child.tag not in (qn("w:rPr"), qn("w:sectPr")):
                if dst_pPr.find(child.tag) is None:
                    dst_pPr.append(deepcopy(child))

def analyze_template(template_path: str | Path, doc: Optional[Document] = None) -> dict:
    """解析模板，返回蓝图：
    {
      "blocks": [段落块（表格内段落 kind="cell"）],
      "tables": [{"index", "kind": "form|data", "headers", "rows", "slots", "table"}],
      "para_slots": [正文「标签+空段」填写位],
      "comments": {归一化文本: 批注},
    }
    """
    doc = doc or Document(str(template_path))
    comments_by_id = _load_comments(template_path)
    body = doc.element.body
    blocks: list[dict] = []
    tables: list[dict] = []
    idx = 0
    tbl_idx = 0

    def walk(element):
        nonlocal idx, tbl_idx
        for child in element.iterchildren():
            tag = child.tag
            if tag == qn("w:p"):
                idx += 1
                para = Paragraph(child, doc)
                text = _para_text(para)
                style = _para_style(para)
                level = _heading_level(para)
                is_heading = level > 0 or _looks_like_title(para)
                blocks.append({
                    "kind": "paragraph", "index": idx, "text": text,
                    "style": style, "level": level, "is_heading": is_heading,
                    "is_empty": _is_empty_para(para), "placeholder": None,
                    "para": para,
                    "comment": _para_comment(para, comments_by_id),
                })
            elif tag == qn("w:tbl"):
                from docx.table import Table
                table = Table(child, doc)
                tbl_idx += 1
                is_form = _table_is_form(table)
                headers, rows = [], []
                for r_i, row in enumerate(table.rows):
                    cells = [c.text.strip() for c in row.cells]
                    if r_i == 0:
                        headers = cells
                    else:
                        rows.append(cells)
                tinfo = {"index": tbl_idx, "kind": "form" if is_form else "data",
                         "headers": headers, "rows": rows, "slots": [], "table": table}
                tables.append(tinfo)
                # 表格内段落也登记（kind="cell"；合并单元格按一份计，保证段落序号与文档顺序一致）
                seen_tc: set[int] = set()
                for r_i, row in enumerate(table.rows):
                    for c_i, cell in enumerate(row.cells):
                        if id(cell._tc) in seen_tc:
                            continue
                        seen_tc.add(id(cell._tc))
                        for para in cell.paragraphs:
                            idx += 1
                            blocks.append({
                                "kind": "cell", "index": idx, "text": _para_text(para),
                                "style": _para_style(para), "level": 0,
                                "is_heading": False,
                                "is_empty": _is_empty_para(para),
                                "placeholder": None, "para": para,
                                "cell": (tbl_idx, r_i, c_i),
                                "comment": _para_comment(para, comments_by_id),
                            })
                walk(child)

    walk(body)

    # 所有表格的填写位（cell 块已全部登记）
    for _t in tables:
        if _t["kind"] == "form":
            _t["slots"] = _extract_form_slots(_t["table"], comments_by_id)

    # 正文「标签+空段」填写位
    para_slots = _extract_para_slots(blocks)
    paras = [b["para"] for b in blocks if b["kind"] == "paragraph"]
    for s in para_slots:
        s["paras"] = paras
        s["para_index"] = 0
        for pi, p in enumerate(paras):
            if re.sub(r"\s+", "", _para_text(p)) == re.sub(r"\s+", "", s["label"]):
                s["para_index"] = pi
                break

    # 占位符识别（{{xxx}} / [xxx] / （xxx） 等，排除正文句子）
    for b in blocks:
        text = b["text"]
        for m in _PLACEHOLDER_RE.finditer(text):
            name = m.group(1).strip()
            if 1 <= len(name) <= 30 and not re.search(r"[，。；！？、]", name):
                b["placeholder"] = name
                b["is_heading"] = False

    return {"blocks": blocks, "tables": tables, "para_slots": para_slots,
            "comments": comments_by_id}
def blueprint_text(info: dict) -> str:
    """把模板蓝图格式化为给 AI 的文本（带编号，便于定位）。"""
    lines = []
    lines.append("【文档结构】")
    for b in info["blocks"]:
        if b["kind"] == "cell":
            continue
        mark = "【标题】" if b.get("is_heading") else ("【空位】" if b.get("is_empty") else "       ")
        ph = f" 占位符:{b['placeholder']}" if b.get("placeholder") else ""
        text = b["text"].strip() if b["text"] else "（空）"
        if text and len(text) > 400:
            text = text[:400] + "……"
        lines.append(f"[{b['index']}] {mark} {text}{ph}")
    for t in info.get("tables", []):
        if t["kind"] == "data":
            hdr = "｜".join(t["headers"][:6])
            n = len(t["rows"])
            lines.append(f"[表{t['index']}] 数据表格 表头: {hdr} 已有数据行: {n}")
    lines.append("【表格填写位】（label 为定位原文；批注为模板要求，必须遵守；信息栏仅在用户明确要求时填写）")
    for t in info.get("tables", []):
        if t["kind"] != "form":
            continue
        for s in t["slots"]:
            tag = "信息栏" if s.get("info_field") else "内容栏"
            cm = f"｜批注:{s['comment']}" if s.get("comment") else ""
            hint = f"｜提示:{s['hint']}" if s.get("hint") else ""
            lines.append(f"[表{t['index']}] {tag}「{s['label']}」{hint}{cm}")
    for s in info.get("para_slots", []):
        tag = "信息栏" if s.get("info_field") else "内容栏"
        lines.append(f"[正文] {tag}「{s['label']}」")

    # 模板正文全文：用户写在 Word 里的要求/说明/注意事项必须传给 AI。
    # 之前这里只截断 60 字且完全跳过表格单元格，导致模板内要求从未被读取。
    lines.append("")
    lines.append("【模板正文全文】（模板里写的每段文字与每个表格单元格。标了【要求】的是包含要求/注意/格式等关键词的语句，必须遵守：格式说明、内容范围、字数要求等；填写内容要贴合这些要求）")
    for b in info["blocks"]:
        t = b["text"].strip()
        if not t:
            continue
        tag = "【要求】" if _looks_like_requirement(t) else ""
        if len(t) > 400:
            t = t[:400] + "……"
        where = "表格" if b["kind"] == "cell" else "正文"
        lines.append(f"[{where}#{b['index']}] {tag}{t}")
    return "\n".join(lines)


def _looks_like_requirement(text: str) -> bool:
    """判断一段模板文本是否像「要求/注意事项」：含要求类关键词，或位于表格的说明单元格。"""
    t = text.strip()
    if not t:
        return False
    kws = ("要求", "注意", "必须", "格式", "不得", "禁止", "请填写", "填写要求", "说明", "提示", "需满足", "应符合", "应包含", "字数", "行距", "字体", "页边距")
    return any(k in t for k in kws)
def _pick_format_anchor(blocks: list[dict], target_index: int = 0,
                        anchor_para: Optional[Paragraph] = None) -> Optional[Paragraph]:
    """选择格式样本：优先锚点之后最近的正文段，其次之前，再退化为任意正文段。

    模板中章节标题后的正文往往是空段落（待填），空段落也可能已带格式
    （如宋体+缩进），因此空段落同样可作为格式样本。
    """
    def _is_body(b: dict) -> bool:
        return b["kind"] == "paragraph" and not b.get("is_heading")

    found = False
    for b in blocks:
        if _is_body(b) and anchor_para is not None and b["para"] is anchor_para:
            found = True
            continue
        if found and _is_body(b):
            return b["para"]
    # 未找到锚点之后：退化为之前的正文段
    for b in reversed(blocks):
        if _is_body(b) and (anchor_para is None or b["para"] is not anchor_para):
            if b["text"].strip() or True:
                return b["para"]
    return None

def _copy_para_format(template_para: Paragraph, new_para: Paragraph) -> None:
    """把模板段落的格式（rPr/pPr）复制到新段落。"""
    src_p = template_para._p
    dst_p = new_para._p
    # 复制 pPr（段落格式：缩进/行距/对齐/边框等），排除编号与大纲级别
    src_ppr = src_p.pPr
    if src_ppr is not None:
        dst_ppr = dst_p.get_or_add_pPr()
        for child in src_ppr:
            tag = child.tag
            if tag in (qn("w:numPr"), qn("w:outlineLvl"), qn("w:sectPr")):
                continue
            dst_ppr.append(deepcopy(child))
    # 复制第一个 run 的 rPr（字体/字号/颜色）
    src_runs = template_para.runs
    if src_runs:
        src_rpr = src_runs[0]._element.rPr
        if src_rpr is not None and new_para.runs:
            dst_rpr = new_para.runs[0]._element.get_or_add_rPr()
            for child in src_rpr:
                if dst_rpr.find(child.tag) is None:
                    dst_rpr.append(deepcopy(child))

def _clone_paragraph_after(doc: Document, anchor_p: Paragraph, text: str, blocks: list[dict]) -> Paragraph:
    """在锚点段落后克隆一个同格式段落并填入文本。"""
    new_p = doc.add_paragraph()
    anchor = _pick_format_anchor(blocks, 0) or anchor_p
    # 用锚点段落做格式样本
    fmt_para = anchor_p if _para_text(anchor_p).strip() else (anchor or anchor_p)
    _copy_para_format(fmt_para, new_p)
    run = new_p.add_run(text)
    # 若格式样本没有 run 格式，复制其 rPr
    if fmt_para.runs and fmt_para.runs[0]._element.rPr is not None:
        rpr = fmt_para.runs[0]._element.rPr
        new_rpr = run._element.get_or_add_rPr()
        for child in rpr:
            if new_rpr.find(child.tag) is None:
                new_rpr.append(deepcopy(child))
    # 移到锚点之后
    anchor_p._p.addnext(new_p._p)
    return new_p

# ---------- 落笔 ----------

def apply_fill(template_path: str | Path, fill: dict,
               log: Callable[[str], None] = print,
               report: Optional[dict] = None) -> list[str]:
    """把 AI 的填写计划写入模板（格式继承模板）。

    fill = {
      "cells":    [{"label": "实验目的及要求", "value": "……"}],  # 表格/正文填写位
      "sections": [{"anchor": "实验目的", "content": "正文……"}],  # 章节正文
      "tables":   [{"headers": [...], "rows": [[...]]}],          # 仅数据表格
    }
    report：可选 dict，就地累计「成功/跳过」统计（成功数 + 跳过明细），供界面展示完成度。
    """
    if report is not None:
        report.setdefault("cells_ok", 0)
        report.setdefault("cells_skip", [])
        report.setdefault("sections_ok", 0)
        report.setdefault("sections_skip", [])
        report.setdefault("tables_ok", 0)
        report.setdefault("tables_skip", [])
    doc = Document(str(template_path))
    info = analyze_template(template_path, doc=doc)
    blocks = info["blocks"]
    applied: list[str] = []
    filled_labels: set[str] = set()
    used_slots: set[int] = set()

    # 1) 表格/正文填写位：按 label 定位写入。
    #    同名位置（如两张表的「实验项目」）：AI 给了多条值则按序对应；
    #    AI 只给一条值则自动填满所有同名位置（同名栏目内容通常一致）。
    cells_by_label: dict[str, list[str]] = {}
    cells_order: list[str] = []
    for cell in fill.get("cells") or []:
        label = str(cell.get("label") or cell.get("placeholder") or "").strip()
        value = str(cell.get("value") or cell.get("content") or "").strip()
        if not label or not value:
            continue
        if label not in cells_by_label:
            cells_by_label[label] = []
            cells_order.append(label)
        cells_by_label[label].append(value)
    for label in cells_order:
        values = cells_by_label[label]
        slots = _find_fill_slots(info, label)
        slots = [s for s in slots if s.get("slot_key") not in used_slots]
        if not slots:
            if not re.search(r"[\[\{（(]", label):
                log(f"   ⚠ 未找到填写位「{label[:30]}」，已跳过")
                if report is not None:
                    report["cells_skip"].append(f"{label[:40]}（模板中无此填写位）")
            continue
        if len(values) >= len(slots):
            plan_pairs = list(zip(slots, values))
        else:
            # 同名栏目位置比 AI 给的值多：前 N 个按序对应，其余用最后一条补齐，避免留空
            plan_pairs = list(zip(slots, values))
            if slots:
                for s in slots[len(values):]:
                    plan_pairs.append((s, values[-1]))
                    if report is not None:
                        report.setdefault("cells_note", []).append(
                            f"「{s.get('label', label)[:30]}」同名位置不足，用最后一条补齐"
                        )
        for slot, value in plan_pairs:
            n = _fill_slot(doc, slot, value)
            if n:
                applied.append(f"填写「{slot['label'][:30]}」（{n} 段）")
                used_slots.add(slot.get("slot_key"))
                filled_labels.add(re.sub(r"\s+", "", slot["label"]))
                if report is not None:
                    report["cells_ok"] += 1
                    # 内容过短的填写位给出提示，方便用户定位
                    if len(str(value).strip()) < 30:
                        report.setdefault("cells_thin", []).append(
                            f"「{slot.get('label', label)[:30]}」内容仅 {len(str(value).strip())} 字，偏短"
                        )
            elif report is not None:
                report["cells_skip"].append(f"{slot.get('label', label)[:40]}（位置已有内容或未匹配）")

    # 2) 占位符替换（{{xxx}} / [xxx]，仅改 run.text 保留格式）
    for cell in fill.get("cells") or []:
        name = str(cell.get("placeholder", "")).strip()
        value = str(cell.get("value", "")).strip()
        if not name or not value:
            continue
        replaced = _replace_placeholder(doc, name, value)
        if replaced:
            applied.append(f"替换占位符「{name}」")

    # 3) 章节填写：anchor 定位标题，正文插到标题后（已按填写位处理的不再重复）
    for sec in fill.get("sections") or []:
        anchor_text = str(sec.get("anchor", "")).strip()
        content = str(sec.get("content", "")).strip()
        if not anchor_text or not content:
            continue
        if re.sub(r"\s+", "", anchor_text) in filled_labels:
            continue
        para = _find_anchor(doc, blocks, anchor_text)
        if para is None:
            log(f"   ⚠ 未找到章节「{anchor_text[:30]}」，已跳过")
            if report is not None:
                report["sections_skip"].append(f"{anchor_text[:40]}（模板中未找到该章节）")
            continue
        n = _insert_after(doc, para, content, blocks)
        if n:
            applied.append(f"填写章节「{anchor_text[:30]}」（{n} 段）")
            if report is not None:
                report["sections_ok"] += 1
        elif report is not None:
            report["sections_skip"].append(f"{anchor_text[:40]}（该章节已有正文，未覆盖）")

    # 4) 数据表格追加行（仅 kind=="data"，表单型表格禁止追加）
    for t in fill.get("tables") or []:
        headers = [str(h).strip() for h in (t.get("headers") or [])]
        rows = [[str(v).strip() for v in (r or [])] for r in (t.get("rows") or [])]
        if not headers or not rows:
            continue
        table = _find_table(info, headers)
        if table is None:
            log(f"   ⚠ 未找到可追加的数据表格「{'、'.join(headers)[:40]}」，已跳过")
            if report is not None:
                report["tables_skip"].append(f"{'、'.join(headers)[:40]}（模板中无匹配的数据表格）")
            continue
        n = _append_rows(table, rows)
        if n:
            applied.append(f"填写表格「{'、'.join(headers)[:30]}」（{n} 行）")
            if report is not None:
                report["tables_ok"] += n
        elif report is not None:
            report["tables_skip"].append(f"{'、'.join(headers)[:40]}（表格行数据为空）")

    doc.save(str(template_path))
    return applied
def _replace_placeholder(doc: Document, name: str, value: str) -> bool:
    """全文（含表格/页眉页脚）替换占位符，只改文本保留格式。"""
    count = 0
    body = doc._element.body
    for t_elem in body.iter(qn("w:t")):
        if t_elem.text:
            new_text = _replace_in_text(t_elem.text, name, value)
            if new_text != t_elem.text:
                t_elem.text = new_text
                count += 1
    return count > 0

_PH_FULL_RE = re.compile(r"[\[\{（(]\s*([^\]\}）)]{1,30}?)\s*[\]\}）)]")

def _replace_in_text(text: str, name: str, value: str) -> str:
    # 精确占位符 {{name}} / [name] / （name）
    pattern = re.compile(
        r"[\[\{（(]\s*" + re.escape(name) + r"\s*[\]\}）)]"
    )
    return pattern.sub(value, text)

def _find_anchor(doc: Document, blocks: list[dict], text: str) -> Optional[Paragraph]:
    norm = re.sub(r"\s+", "", text)
    best = None
    best_score = 0
    for b in blocks:
        if b["kind"] != "paragraph":
            continue
        t = re.sub(r"\s+", "", b["text"])
        if not t:
            continue
        if t == norm:
            return b["para"]
        if norm in t or t in norm:
            score = min(len(norm), len(t))
            if score > best_score:
                best_score = score
                best = b["para"]
    return best if best_score >= 2 else None

def _insert_after(doc: Document, anchor_para: Paragraph, content: str,
                  blocks: list[dict]) -> int:
    """把多行内容填入锚点标题之后的正文区。

    优先【替换】锚点后紧跟的占位/示例段落（空段、【示例】、下划线等），
    无占位段时再新增段落；格式继承模板正文。
    """
    lines = [_clean_fill_line(ln) for ln in content.splitlines()]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    if not lines:
        return 0
    # 找锚点后紧跟的正文段落（可能有多个，取第一个）
    target = _next_body_para(anchor_para, blocks)
    if target is not None and _is_placeholder_para(target):
        # 替换占位段落：清空原文，写入新内容，保留格式
        first_line = lines[0]
        _replace_para_text_keep_format(target, first_line)
        anchor_para = target
        rest = lines[1:]
        for line in rest:
            if not line.strip():
                continue
            new_p = doc.add_paragraph()
            _copy_para_format(target, new_p)
            if target.runs and target.runs[0]._element.rPr is not None:
                rpr = target.runs[0]._element.rPr
                run = new_p.add_run(line)
                new_rpr = run._element.get_or_add_rPr()
                for child in rpr:
                    if new_rpr.find(child.tag) is None:
                        new_rpr.append(deepcopy(child))
            anchor_para._p.addnext(new_p._p)
            anchor_para = new_p
        return len([l for l in lines if l.strip()])
    # 无占位段：直接新增
    sample = _pick_format_anchor(blocks, anchor_para=anchor_para) or _pick_after_anchor(blocks, anchor_para)
    for line in lines:
        if not line.strip():
            continue
        new_p = doc.add_paragraph()
        if sample is not None:
            _copy_para_format(sample, new_p)
            if sample.runs and sample.runs[0]._element.rPr is not None:
                rpr = sample.runs[0]._element.rPr
                run = new_p.add_run(line)
                new_rpr = run._element.get_or_add_rPr()
                for child in rpr:
                    if new_rpr.find(child.tag) is None:
                        new_rpr.append(deepcopy(child))
        else:
            new_p.add_run(line)
        anchor_para._p.addnext(new_p._p)
        anchor_para = new_p
    return len([l for l in lines if l.strip()])


def _next_body_para(anchor_para: Paragraph, blocks: list[dict]):
    """锚点标题之后紧邻的第一个正文段落。"""
    found = False
    for b in blocks:
        if b["kind"] != "paragraph":
            continue
        if b["para"] is anchor_para:
            found = True
            continue
        if found and not b.get("is_heading"):
            return b["para"]
    return None


def _is_placeholder_para(para: Paragraph) -> bool:
    """判断段落是否为待填占位：空段 / 【示例】 / 下划线 / 提示语。"""
    text = _para_text(para).strip()
    if not text:
        return True
    return any(p.match(text) for p in _EMPTY_PATTERNS)


def _replace_para_text_keep_format(para: Paragraph, text: str) -> None:
    """清空段落所有 run 文本，用第一个 run（保留格式）写入新文本。"""
    runs = para.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        para.add_run(text)

def _pick_after_anchor(blocks: list[dict], anchor_para: Paragraph):
    """找锚点之后的第一个正文段作为格式样本。"""
    found = False
    for b in blocks:
        if b["kind"] != "paragraph" or b.get("is_heading") or b.get("is_empty"):
            continue
        if b["para"] is anchor_para:
            found = True
            continue
        if found and b["text"].strip():
            return b["para"]
    return None

def _find_table(info: dict, headers: list[str]):
    """按表头序列匹配【数据型】表格（表单型禁止追加行）。"""
    clean = [re.sub(r"\s+", "", h) for h in headers if h]
    if not clean:
        return None
    joined = "".join(clean)
    best, best_score = None, 0
    for t in info.get("tables", []):
        if t["kind"] != "data":
            continue
        cells = [re.sub(r"\s+", "", h) for h in t["headers"]]
        if "".join(cells) == joined:
            return t["table"]
        score = sum(1 for c in cells if c and any(c == h or c in h or h in c for h in clean))
        if score > best_score:
            best_score, best = score, t["table"]
    return best if best_score >= max(1, len(clean) * 0.6) else None
def _append_rows(table, rows: list[list[str]]) -> int:
    """追加数据行：复制模板最后一行（或表头行）的格式。"""
    template_row = None
    if len(table.rows) >= 2:
        template_row = table.rows[1]
    elif len(table.rows) == 1:
        template_row = table.rows[0]
    added = 0
    for row in rows:
        values = [_clean_fill_text(v) for v in row]
        if not any(v.strip() for v in values):
            continue
        new_row = table.add_row()
        for j, val in enumerate(values):
            if j >= len(new_row.cells):
                break
            cell = new_row.cells[j]
            p = cell.paragraphs[0]
            if p.runs:
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
            run = p.add_run(val.strip())
            if template_row is not None:
                try:
                    src = template_row.cells[j].paragraphs[0]
                    if src.runs and src.runs[0]._element.rPr is not None:
                        run._element.insert(0, deepcopy(src.runs[0]._element.rPr))
                except Exception:
                    pass
        added += 1
    return added

__all__ = ["analyze_template", "blueprint_text", "apply_fill"]
