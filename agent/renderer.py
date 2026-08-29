"""Markdown 正文 -> 排版精美的 .docx（基于 python-docx）。

字体一致性设计：
- 样式级统一：Normal / Heading 1-4 / 列表样式显式设置字体，杜绝回退 Calibri/等线
- 中英文字体分离：模板提供 latin 字体（如 Times New Roman）与中文 eastAsia 字体（如 宋体）时分开设置，
  保持中文正式文档「汉字宋体 + 西文 Times New Roman」的规范观感
- 文本清理：清理未闭合的 ** / * / ` 标记与 LaTeX 数学残留（\\( \\) \\frac 等），避免乱码
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------- 排版预设 ----------
STYLE_PRESETS: dict[str, dict] = {
    "business": {"heading_color": "1F4E79", "body_font": "微软雅黑", "code_font": "Consolas"},
    "report":   {"heading_color": "2E74B5", "body_font": "微软雅黑", "code_font": "Consolas"},
    "academic": {"heading_color": "000000", "body_font": "宋体", "code_font": "Consolas"},
    "creative": {"heading_color": "C55A11", "body_font": "微软雅黑", "code_font": "Consolas"},
    "default":  {"heading_color": "1F3864", "body_font": "微软雅黑", "code_font": "Consolas"},
}

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+`)")
CODE_FENCE_RE = re.compile(r"^\s*```")
TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|[\s:|-]+\|\s*$")


def _is_table_sep(line: str) -> bool:
    """分隔行判定：首尾竖线标准格式，或宽松格式（单元格仅由 - : 空格组成）。"""
    s = line.strip()
    if TABLE_SEP_RE.match(s):
        return True
    if s.count("|") < 2:
        return False
    cells = [c.strip() for c in s.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{1,}:?", c) for c in cells)


def _is_table_row(line: str) -> bool:
    """表格行判定：首尾竖线（标准 markdown）或行内含 >=2 个竖线（宽松格式）。"""
    s = line.strip()
    if TABLE_ROW_RE.match(s):
        return True
    return s.count("|") >= 2
BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
NUMBER_RE = re.compile(r"^(\s*)\d+[.、)]\s+(.*)$")
HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")
QUOTE_RE = re.compile(r"^>\s?(.*)$")
IMAGE_RE = re.compile(r"^!\[(.*?)\]\(([^)]+)\)\s*$")

# 轻量 LaTeX 数学残留清理（按序执行）
LATEX_CLEANUPS = [
    (re.compile(r"\\\(|\\\)|\\\[|\\\]"), ""),
    (re.compile(r"\$\$([^$]*)\$\$"), lambda m: m.group(1)),
    (re.compile(r"\$([^$]*)\$"), lambda m: m.group(1)),
    (re.compile(r"\\frac\{([^{}]*)\}\{([^{}]*)\}"), lambda m: f"{m.group(1)}/{m.group(2)}"),
    (re.compile(r"\\bar\{([^{}]*)\}"), lambda m: m.group(1) + "\u0304"),
    (re.compile(r"\\rho"), "ρ"),
    (re.compile(r"\\Omega"), "Ω"),
    (re.compile(r"\\times"), "×"),
    (re.compile(r"\\cdot"), "·"),
    (re.compile(r"\\left|\\right"), ""),
    (re.compile(r"_\{([^{}]*)\}"), lambda m: m.group(1)),
    (re.compile(r"_([0-9A-Za-z\u0370-\u03ff])"), lambda m: m.group(1)),
    (re.compile(r"\^\{([^{}]*)\}"), lambda m: m.group(1)),
    (re.compile(r"\^([0-9A-Za-z\u0370-\u03ff])"), lambda m: m.group(1)),
    (re.compile(r"\\[a-zA-Z]+"), ""),
    (re.compile(r"[\{\}]"), ""),
]


def _clean_latex(text: str) -> str:
    for pat, rep in LATEX_CLEANUPS:
        text = pat.sub(rep, text)
    return text


def _clean_markdown_text(text: str) -> str:
    """清理标题/纯文本中的残留 Markdown 标记与 LaTeX。"""
    text = _clean_latex(text)
    text = text.replace("**", "").replace("`", "").strip()
    return text


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def _set_run_font(run, name: str, size: float, bold: bool = False, color: Optional[str] = None,
                  italic: bool = False, latin: Optional[str] = None) -> None:
    """设置 run 字体：name 为中文（eastAsia）字体，latin 为西文/数字字体（缺省同 name）。"""
    latin = latin or name
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = _rgb(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), name)


def _body_font(preset: dict) -> str:
    return preset.get("body_font") or "微软雅黑"


def _latin_font(preset: dict) -> Optional[str]:
    return preset.get("body_font_latin") or None


def _heading_font(preset: dict) -> str:
    return preset.get("heading_font") or _body_font(preset)


def _set_body_run(run, preset: dict, size: float, bold: bool = False,
                  color: Optional[str] = None, italic: bool = False) -> None:
    """按预设设置正文 run（自动应用中英文字体分离）。"""
    _set_run_font(run, _body_font(preset), size, bold=bold, color=color, italic=italic,
                  latin=_latin_font(preset))


def _set_style_font(style, name: str, size: float, bold: bool = False,
                    color: Optional[str] = None, latin: Optional[str] = None) -> None:
    """设置样式级字体（Normal/Heading/列表），避免 Word 回退默认字体。"""
    latin = latin or name
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = _rgb(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), name)


def _apply_line_spacing(pf, preset: dict) -> None:
    """按标准设置行距：multiple=倍数，exact=固定磅值。"""
    rule = preset.get("line_spacing_rule") or "multiple"
    value = preset.get("line_spacing") or 1.5
    if rule == "exact":
        pf.line_spacing = Pt(float(value))
    else:
        pf.line_spacing = float(value)


def _setup_styles(doc: Document, preset: dict) -> None:
    """统一 Normal/标题/列表样式字体，与 run 级设置一致，杜绝回退。"""
    body = _body_font(preset)
    latin = _latin_font(preset)
    heading = _heading_font(preset)
    body_size = preset.get("body_size") or 12.0
    heading_color = preset.get("heading_color") or "1F3864"
    heading_fonts = preset.get("heading_fonts") or {}
    heading_bold = preset.get("heading_bold") or {}
    sizes = {1: 16, 2: 14, 3: 12.5, 4: 12}
    hs = preset.get("heading_sizes") or {}
    for _lv, _sz in hs.items():
        sizes[int(_lv)] = float(_sz)

    try:
        _set_style_font(doc.styles["Normal"], body, body_size, latin=latin)
        _apply_line_spacing(doc.styles["Normal"].paragraph_format, preset)
        doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    except KeyError:
        pass
    for lv in (1, 2, 3, 4):
        try:
            st = doc.styles[f"Heading {lv}"]
        except KeyError:
            continue
        font = heading_fonts.get(lv) or heading
        bold = heading_bold.get(lv, True)
        _set_style_font(st, font, sizes.get(lv, 12), bold=bold, color=heading_color, latin=latin)
        st.paragraph_format.space_before = Pt(14 if lv <= 2 else 10)
        st.paragraph_format.space_after = Pt(8)
    for name in ("List Bullet", "List Bullet 2", "List Bullet 3",
                 "List Number", "List Number 2", "List Number 3"):
        try:
            _set_style_font(doc.styles[name], body, body_size, latin=latin)
        except KeyError:
            continue


def _shade_paragraph(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def _shade_cell(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def _setup_document(doc: Document, page: Optional[dict] = None) -> None:
    for section in doc.sections:
        if page:  # 模板驱动：完全按模板页面设置
            section.page_width = Cm(page.get("width_cm", 21.0))
            section.page_height = Cm(page.get("height_cm", 29.7))
            section.top_margin = Cm(page.get("top_cm", 2.5))
            section.bottom_margin = Cm(page.get("bottom_cm", 2.5))
            section.left_margin = Cm(page.get("left_cm", 2.8))
            section.right_margin = Cm(page.get("right_cm", 2.8))
        else:
            section.page_width = Cm(21.0)   # A4
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.8)
            section.right_margin = Cm(2.8)


def _add_footer_page_number(doc: Document) -> None:
    footer = doc.sections[0].footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("第 ")
    _set_run_font(run, "微软雅黑", 9, color="808080")
    run = para.add_run()
    _set_run_font(run, "微软雅黑", 9, color="808080")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run = para.add_run(" 页")
    _set_run_font(run, "微软雅黑", 9, color="808080")


def _add_title_page(doc: Document, title: str, preset: dict, date_text: str) -> None:
    # 顶部留白（两行即可，保持版面干净）
    for _ in range(2):
        doc.add_paragraph()
    title_font = preset.get("title_font") or _body_font(preset)
    title_size = preset.get("title_size") or 26
    title_bold = preset.get("title_bold", True)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    _set_run_font(run, title_font, title_size, bold=title_bold,
                  color=preset.get("heading_color") or "1F3864", latin=_latin_font(preset))
    para.paragraph_format.space_after = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(date_text)
    _set_body_run(run, preset, 12, color="595959")

    doc.add_page_break()


def _enable_field_update(doc: Document) -> None:
    """开启 w:updateFields：Word 打开文档时自动刷新目录等域。"""
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        upd = OxmlElement("w:updateFields")
        upd.set(qn("w:val"), "true")
        settings.append(upd)


def _add_toc_field(doc: Document, preset: dict) -> None:
    # 目录标题用普通段落（outlineLvl=body text），否则 Word 会把「目录」自身收进目录
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("目录")
    heading_font = _heading_font(preset)
    sizes = preset.get("heading_sizes") or {}
    _set_run_font(run, heading_font, sizes.get(1, 16), bold=True,
                  color=preset.get("heading_color") or "1F3864", latin=_latin_font(preset))
    para.paragraph_format.space_after = Pt(12)
    ppr = para._p.get_or_add_pPr()
    olvl = OxmlElement("w:outlineLvl")
    olvl.set(qn("w:val"), "9")
    ppr.append(olvl)
    para = doc.add_paragraph()
    run = para.add_run()
    _set_run_font(run, "微软雅黑", 12)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "（目录：打开文档时将自动更新）"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)
    doc.add_page_break()


def _add_heading(doc: Document, level: int, text: str, preset: dict) -> None:
    # 使用原生 Heading 样式（目录/大纲可用），并统一字体
    sizes = {1: 16, 2: 14, 3: 12.5, 4: 12}
    hs = preset.get("heading_sizes") or {}
    for _lv, _sz in hs.items():
        sizes[int(_lv)] = float(_sz)
    clean_text = _clean_markdown_text(text)
    para = doc.add_heading(clean_text or "（空标题）", level=min(level, 4))
    para.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    para.paragraph_format.space_after = Pt(8)
    if level <= 2:
        para.paragraph_format.keep_with_next = True
    heading_fonts = preset.get("heading_fonts") or {}
    heading_bold = preset.get("heading_bold") or {}
    font = heading_fonts.get(level) or _heading_font(preset)
    bold = heading_bold.get(level, True)
    for run in para.runs:
        _set_run_font(run, font, sizes.get(level, 12), bold=bold,
                      color=preset.get("heading_color") or "1F3864", latin=_latin_font(preset))
    if level == 1:  # 一级标题下加一条分隔线
        ppr = para._p.get_or_add_pPr()
        pbd = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), preset.get("heading_color") or "1F3864")
        pbd.append(bottom)
        ppr.append(pbd)


def _add_body(doc: Document, text: str, preset: dict) -> None:
    para = doc.add_paragraph()
    body_size = preset.get("body_size") or 12.0
    indent_chars = preset.get("first_line_indent_chars")
    if indent_chars:
        para.paragraph_format.first_line_indent = Pt(body_size * float(indent_chars))
    else:
        para.paragraph_format.first_line_indent = Pt(0)
    _apply_line_spacing(para.paragraph_format, preset)
    if preset.get("justify", True):
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 正式文档两端对齐
    _add_inline_runs(para, text, preset)


def _add_inline_runs(para, text: str, preset: dict) -> None:
    """解析 **加粗**、*斜体*、`代码` 等行内格式；清理残留标记与 LaTeX。"""
    body_size = preset.get("body_size") or 12.0
    text = _clean_latex(text)
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            run = para.add_run(_clean_markdown_text(token[2:-2]))
            _set_body_run(run, preset, body_size, bold=True)
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = para.add_run(_clean_markdown_text(token[1:-1]))
            _set_body_run(run, preset, body_size, italic=True)
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            run = para.add_run(token[1:-1])
            _set_run_font(run, preset.get("code_font") or "Consolas", 11, color="C7254E")
        else:
            run = para.add_run(_clean_markdown_text(token))
            _set_body_run(run, preset, body_size)


def _add_list_item(doc: Document, text: str, ordered: bool, indent_level: int, preset: dict,
                   list_seq: int | None = None) -> None:
    """List item. Ordered lists use manual numbering (restarts at 1 per list)
    to avoid Word sharing one numbering sequence across separate lists."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.4
    indent_cm = 0.74 * (indent_level + 1)
    para.paragraph_format.left_indent = Cm(indent_cm)
    para.paragraph_format.first_line_indent = Cm(-0.74)
    if ordered:
        seq = list_seq if list_seq is not None else 1
        marker = para.add_run(f"{seq}. ")
        _set_body_run(marker, preset, preset.get("body_size") or 12.0, bold=True)
    else:
        marker = para.add_run("\u2022 ")
        _set_body_run(marker, preset, preset.get("body_size") or 12.0, bold=True)
    _add_inline_runs(para, text, preset)


def _add_code_block(doc: Document, lines: list[str], preset: dict) -> None:
    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.left_indent = Cm(0.6)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(line if line else " ")
        _set_run_font(run, preset.get("code_font") or "Consolas", 10, color="333333")
        _shade_paragraph(para, "F2F2F2")


def _add_blockquote(doc: Document, text: str, preset: dict) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.8)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(_clean_markdown_text(text))
    _set_body_run(run, preset, 11.5, italic=True, color="595959")
    ppr = para._p.get_or_add_pPr()
    pbd = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), preset.get("heading_color") or "1F3864")
    pbd.append(left)
    ppr.append(pbd)


def _add_image(doc: Document, path: str, caption: str, preset: dict,
                 extra_images: Optional[list] = None) -> None:
    """插入图片（居中 + 图注）。路径不存在时按文件名在 extra_images 中兜底匹配。"""
    img_path = Path(path)
    if not img_path.is_file() and extra_images:
        for cand in extra_images:
            if Path(cand).name == img_path.name:
                img_path = Path(cand)
                break
    try:
        # 图片宽度：默认不超过版心 13.5cm；有像素信息时按 96dpi 折算，避免过小
        width_cm = 13.5
        try:
            from PIL import Image as _PILImage
            with _PILImage.open(str(img_path)) as im:
                w_px, h_px = im.size
            if w_px > 0:
                width_cm = min(13.5, w_px * 2.54 / 96.0)
        except Exception:
            pass
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(img_path), width=Cm(width_cm))
        para.paragraph_format.space_after = Pt(2)
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c_run = cap.add_run("图：%s" % _clean_markdown_text(caption))
            _set_body_run(c_run, preset, 9, color="808080")
            cap.paragraph_format.space_after = Pt(10)
    except Exception:
        para = doc.add_paragraph()
        run = para.add_run("[图片未找到：%s]" % path)
        _set_body_run(run, preset, 10.5, color="BFBFBF")


def _is_light_color(hex_color: str) -> bool:
    """按亮度判断颜色深浅（表头底色的文字颜色选择用）。"""
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
    except (ValueError, IndexError):
        return False
    # 相对亮度（加权）：>0.6 视为浅色
    lum = (0.299 * r + 0.587 * g + 0.114 * b) / 255.0
    return lum > 0.6


def _set_table_col_widths(table, rows: list[list[str]], preset: dict) -> None:
    """按内容长度自适应列宽（总宽约 15cm，居中表格内等比例分配）。"""
    ncols = max(len(r) for r in rows)
    widths = [0.0] * ncols
    for row in rows:
        for j in range(ncols):
            text = (row[j] if j < len(row) else "")
            # 中文按 2 字符宽估算；英文/数字按 1
            w = sum(2.0 if ord(ch) > 0x2E7F else 1.0 for ch in text)
            widths[j] = max(widths[j], w)
    total_w = sum(widths) or 1.0
    usable_cm = 15.0
    for j in range(ncols):
        ratio = (widths[j] + 1.0) / (total_w + ncols)
        width_cm = max(1.5, min(6.5, usable_cm * ratio))
        for row in table.rows:
            try:
                row.cells[j].width = Cm(width_cm)
            except Exception:
                pass


def _add_table(doc: Document, rows: list[list[str]], preset: dict) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    head_color = preset.get("heading_color") or "1F3864"
    head_text_color = "333333" if _is_light_color(head_color) else "FFFFFF"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = (row[j] if j < len(row) else "").strip()
            cell = table.cell(i, j)
            cell.paragraphs[0].text = ""
            para = cell.paragraphs[0]
            run = para.add_run(_clean_markdown_text(cell_text))
            if i == 0:
                _set_body_run(run, preset, 11, bold=True, color=head_text_color)
                _shade_cell(cell, head_color)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                _set_body_run(run, preset, 11)
    _set_table_col_widths(table, rows, preset)
    # 表后空行
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


# 标准/模板 preset 可覆盖的字段（优先级：模板 > 标准 > 通用风格）
PRESET_KEYS = (
    "heading_color", "body_font", "body_font_latin", "heading_font", "heading_fonts",
    "heading_bold", "body_size", "heading_sizes", "page", "line_spacing",
    "line_spacing_rule", "first_line_indent_chars", "justify", "title_font",
    "title_size", "title_bold", "title_page", "toc", "code_font",
)


def _merge_preset(base: dict, override: dict) -> dict:
    merged = dict(base)
    for _k in PRESET_KEYS:
        if override.get(_k) is not None:
            merged[_k] = override[_k]
    return merged


def render_markdown_to_docx(
    markdown: str,
    output_path: Path,
    title: str,
    style: str = "default",
    toc: bool = False,
    language: str = "zh-CN",
    extra_images: Optional[list] = None,
    template_preset: Optional[dict] = None,
    standard_preset: Optional[dict] = None,
) -> Path:
    """渲染 Markdown 为 docx 文件，返回输出路径。

    优先级：template_preset（模板文档）> standard_preset（国标/规范）> style 通用预设。
    """
    preset = STYLE_PRESETS.get(style, STYLE_PRESETS["default"])
    if standard_preset:
        preset = _merge_preset(preset, standard_preset)
    if template_preset:
        preset = _merge_preset(preset, template_preset)
    doc = Document()
    _setup_document(doc, preset.get("page"))
    _setup_styles(doc, preset)
    _add_footer_page_number(doc)

    from datetime import datetime
    date_text = datetime.now().strftime("%Y年%m月%d日")

    # 封面与目录开关：模板优先，其次标准，最后取参数
    show_title_page = not template_preset
    if show_title_page and standard_preset and standard_preset.get("title_page") is not None:
        show_title_page = bool(standard_preset["title_page"])
    show_toc = bool(toc)
    if standard_preset and standard_preset.get("toc") is not None:
        show_toc = bool(standard_preset["toc"])
    if template_preset:
        show_toc = False

    if show_title_page:
        _add_title_page(doc, _clean_markdown_text(title), preset, date_text)
    if show_toc:
        _add_toc_field(doc, preset)
        _enable_field_update(doc)

    lines = markdown.splitlines()
    i = 0
    n = len(lines)
    ordered_seq: int | None = None
    while i < n:
        raw = lines[i]
        stripped = raw.strip()
        # 代码围栏
        if CODE_FENCE_RE.match(raw):
            block: list[str] = []
            i += 1
            while i < n and not CODE_FENCE_RE.match(lines[i]):
                block.append(lines[i])
                i += 1
            i += 1  # 跳过结尾围栏
            _add_code_block(doc, block, preset)
            continue
        if not stripped:
            i += 1
            continue
        # 图片
        m_img = IMAGE_RE.match(raw)
        if m_img:
            _add_image(doc, m_img.group(2).strip(), m_img.group(1).strip(), preset, extra_images)
            i += 1
            ordered_seq = None
            continue
        # 标题
        m = HEADING_RE.match(raw)
        if m:
            _add_heading(doc, len(m.group(1)), m.group(2).strip(), preset)
            i += 1
            ordered_seq = None
            continue
        # 表格（支持宽松格式：行内含 >=2 个竖线即可识别）
        if _is_table_row(raw) and i + 1 < n and _is_table_sep(lines[i + 1]):
            header = [c.strip() for c in raw.strip().strip("|").split("|")]
            body = []
            i += 2
            while i < n and _is_table_row(lines[i]):
                body.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            _add_table(doc, [header] + body, preset)
            ordered_seq = None
            continue
        # 引用
        if QUOTE_RE.match(raw):
            _add_blockquote(doc, QUOTE_RE.match(raw).group(1), preset)
            i += 1
            ordered_seq = None
            continue
        # 列表
        mb = BULLET_RE.match(raw)
        mn = NUMBER_RE.match(raw)
        if mb:
            indent = len(mb.group(1)) // 2
            _add_list_item(doc, mb.group(2), ordered=False, indent_level=indent, preset=preset)
            i += 1
            ordered_seq = None
            continue
        if mn:
            indent = len(mn.group(1)) // 2
            if ordered_seq is None:
                ordered_seq = 1
            _add_list_item(doc, mn.group(2), ordered=True, indent_level=indent, preset=preset,
                           list_seq=ordered_seq)
            ordered_seq += 1
            i += 1
            continue
        # 普通段落
        _add_body(doc, stripped, preset)
        i += 1
        ordered_seq = None

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
