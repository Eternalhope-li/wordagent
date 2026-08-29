"""程序化质量校验门禁（本地执行，不消耗 LLM tokens）。

对齐 officecli 的 Verification Gate（validate → issues → outline → 视觉）思路，
在 WordAgent 内实现纯本地版本，生成/编辑完成后自动运行：

- 完整性：文档能否被 python-docx 正常打开
- 标题层级：真实 Heading 样式、无跳级、无空标题
- 空段污染：正文连续空段（标题页刻意留白除外）
- Markdown 残留：**、`、未转换的表格行、图片占位符
- 字体一致性：段落内 run 字体与段落主字体不一致（"字体有粗有细"的根因之一）
- 表格边框：所有表格必须有边框（Table Grid / 显式 tblBorders）
- 内容充实度：每个标题章节正文不得过短（防止"没写后续"的半截内容）
- 图片图注：图片后应有"图：xx"图注

自动修复（安全项）：
- 清理 Markdown 残留标记
- 压缩多余连续空段
- 给无边框表格补上边框
- 归一正文 run 字体到 Normal 样式字体（仅当 run 无显式 eastAsia 或回退 Calibri/等线）
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# 章节最小正文长度（字符）：低于此值视为"没写后续/半截内容"
MIN_SECTION_BODY = 40
# 连续空段阈值：超过该数量视为空段污染（标题页顶部刻意留白除外）
EMPTY_RUN_LIMIT = 2
# 可自动归一化的"回退字体"：Word 未设置 eastAsia 时常见的默认值
FALLBACK_FONTS = {"calibri", "等线", "dengxian", "宋体", "simsun"}


def _iter_body_paragraphs(doc: Document):
    """按文档顺序遍历正文段落（默认仅 body 直接子元素，不进入表格）。"""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)


def _all_paragraphs(doc: Document):
    """遍历正文与表格内所有段落（校验内容充实度时用）。"""
    body = doc.element.body

    def walk(element):
        for child in element.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, doc)
            elif child.tag == qn("w:tbl"):
                yield from walk(child)

    yield from walk(body)


def _heading_level(para: Paragraph) -> int:
    style = para.style
    name = (style.name if style is not None else "") or ""
    if name.lower().startswith("heading"):
        digits = name[7:].strip()
        if digits.isdigit():
            return int(digits)
    base = style.base_style if style is not None else None
    while base is not None:
        bname = (base.name if base.name else "") or ""
        if bname.lower().startswith("heading"):
            digits = bname[7:].strip()
            if digits.isdigit():
                return int(digits)
        base = base.base_style
    ppr = para._p.pPr
    if ppr is not None:
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is not None:
            try:
                val = int(outline.get(qn("w:val"), "0"))
                # outlineLvl >= 9 表示“正文/不参与大纲”（目录标题等），不算标题
                if val >= 9:
                    return 0
                return val + 1
            except ValueError:
                 return 1
    return 0


def _run_east_asia_font(run) -> Optional[str]:
    rpr = run._element.rPr
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return None
    return rfonts.get(qn("w:eastAsia")) or None


def _normal_style_font(doc: Document) -> Optional[str]:
    try:
        st = doc.styles["Normal"]
    except KeyError:
        return None
    rpr = st.element.rPr
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return None
    return rfonts.get(qn("w:eastAsia")) or None


def _table_has_borders(table) -> bool:
    """检查表格是否有边框（Table Grid 样式或显式 tblBorders）。"""
    try:
        style_name = (table.style.name or "").lower()
        if "grid" in style_name or "网格" in style_name:
            return True
    except Exception:
        pass
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return False
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        return False
    for edge in borders:
        if edge.get(qn("w:val")) not in (None, "none", "nil"):
            return True
    return False


def _ensure_table_borders(table) -> None:
    """给表格补上全边框（安全：仅在无边框时执行）。"""
    try:
        table.style = "Table Grid"
        return
    except KeyError:
        pass
    tbl_pr = table._tbl.get_or_add_tblPr()
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn("w:{0}".format(edge_name)))
        if edge is None:
            edge = OxmlElement("w:{0}".format(edge_name))
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "000000")


def _strip_md_residue(text: str) -> str:
    """清理段落文本中的 Markdown 残留：**、`、行首 #、未转换的竖线。"""
    t = text
    t = re.sub(r"\*\*", "", t)
    t = t.replace("`", "")
    t = re.sub(r"(^|\s)#{1,6}\s+", r"\1", t)
    # 未转换的 markdown 表格行：去掉首尾竖线，保留单元格内容
    if t.count("|") >= 2:
        cells = [c.strip() for c in t.strip().strip("|").split("|")]
        t = " | ".join(cells)
    return t.strip()


def _paragraph_has_drawing(para: Paragraph) -> bool:
    return para._p.find(qn("w:drawing")) is not None or para._p.find(qn("w:pict")) is not None


@dataclass
class QualityReport:
    path: str = ""
    issues: list = field(default_factory=list)   # 必须处理的问题
    warnings: list = field(default_factory=list)  # 提示性建议
    fixes: list = field(default_factory=list)     # 已自动修复的项
    ok: bool = True                               # issues 为空即通过

    def log_lines(self):
        lines = []
        if self.fixes:
            lines.append("   🔧 已自动修复 {0} 项：".format(len(self.fixes)))
            for f in self.fixes:
                lines.append("      ✓ {0}".format(f))
        if self.issues:
            lines.append("   ⚠ 质量检查发现 {0} 个问题：".format(len(self.issues)))
            for i in self.issues:
                lines.append("      · {0}".format(i))
        for w in self.warnings:
            lines.append("      · {0}".format(w))
        if not self.issues and not self.warnings:
            lines.append("   ✓ 质量检查通过：标题层级/字体/表格/内容完整度均正常。")
        return lines


def _check_headings(doc: Document, report: QualityReport) -> None:
    levels = []
    for para in _iter_body_paragraphs(doc):
        lvl = _heading_level(para)
        if lvl:
            levels.append((lvl, para.text.strip()))
    if not levels:
        report.warnings.append("文档没有任何标题样式（全为正文段落）；正式报告建议使用 Heading 1/2 层级。")
        return
    has_h1 = any(lvl == 1 for lvl, _ in levels)
    for lvl, text in levels:
        if not text:
            report.issues.append("存在空标题（Heading {0}），标题不能为空。".format(lvl))
        if lvl == 1:
            continue
        if not has_h1:
            report.issues.append("标题「{0}」是 Heading {1}，但文档没有 Heading 1 一级标题。".format(text[:30], lvl))
    prev = levels[0][0]
    for lvl, text in levels[1:]:
        if lvl - prev > 1:
            report.issues.append("标题层级跳级：Heading {0} → Heading {1}（「{2}」）。".format(prev, lvl, text[:30]))
        prev = lvl


def _check_empty_paragraphs(doc: Document, report: QualityReport) -> None:
    paras = list(_iter_body_paragraphs(doc))
    run_start = 0
    # 标题页顶部刻意留白（通常 1-2 个空段）不算污染
    for i, p in enumerate(paras[:4]):
        if not p.text.strip():
            run_start = i + 1
        else:
            break
    empty_run = 0
    for i in range(run_start, len(paras)):
        p = paras[i]
        if not p.text.strip():
            empty_run += 1
        else:
            if empty_run > EMPTY_RUN_LIMIT:
                report.issues.append("第 {0} 段附近有 {1} 个连续空段落（空段污染）。".format(i - empty_run + 1, empty_run))
            empty_run = 0
    if empty_run > EMPTY_RUN_LIMIT:
        report.issues.append("文档末尾有 {0} 个连续空段落。".format(empty_run))


def _check_md_residue(doc: Document, report: QualityReport) -> None:
    count = 0
    for para in _iter_body_paragraphs(doc):
        text = para.text or ""
        if not text.strip():
            continue
        if "**" in text or "`" in text:
            report.issues.append("段落含 Markdown 残留标记：{0}…".format(text[:40]))
            count += 1
        elif re.match(r"^\s*#{1,6}\s", text):
            report.issues.append("段落以 Markdown 标题标记开头：{0}…".format(text[:40]))
            count += 1
        elif text.count("|") >= 2 and _heading_level(para) == 0:
            report.issues.append("段落疑似未转换的表格行：{0}…".format(text[:40]))
            count += 1
        if count >= 6:
            report.issues.append("（更多 Markdown 残留省略……）")
            break


def _check_font_consistency(doc: Document, report: QualityReport) -> None:
    body_font = _normal_style_font(doc)
    for para in _iter_body_paragraphs(doc):
        if _heading_level(para):
            continue
        runs = [r for r in para.runs if (r.text or "").strip()]
        if len(runs) < 2:
            continue
        fonts = set()
        for r in runs:
            f = _run_east_asia_font(r)
            if f is not None:  # 未显式设置的 run 继承 Normal 样式，不构成“混合”
                fonts.add(f.lower())
        if body_font:
            fonts.discard(body_font.lower())
        if len(fonts) > 1:
            shown = " / ".join(sorted(x or "未设置" for x in fonts))
            report.issues.append("段落内字体混合（{0}）：{1}…".format(shown, para.text[:30]))
        if body_font:
            for r in runs:
                f = (_run_east_asia_font(r) or "").lower()
                if f and f != body_font.lower() and f in FALLBACK_FONTS:
                    report.warnings.append("段落字体回退为「{0}」（正文样式应为「{1}」）：{2}…".format(f, body_font, para.text[:24]))


def _check_tables(doc: Document, report: QualityReport) -> None:
    for i, table in enumerate(doc.tables, 1):
        if not _table_has_borders(table):
            report.issues.append("第 {0} 个表格没有边框（Table Grid）。".format(i))


def _check_section_content(doc: Document, report: QualityReport) -> None:
    paras = list(_all_paragraphs(doc))
    current = None
    buf = []
    sections = []

    def flush():
        if current:
            sections.append((current, "".join(buf)))

    for para in paras:
        lvl = _heading_level(para)
        if lvl:
            flush()
            current = para.text.strip() or "(空标题 Heading {0})".format(lvl)
            buf = []
        else:
            buf.append(para.text or "")
    flush()
    for title, body in sections:
        if len(body.strip()) < MIN_SECTION_BODY:
            report.issues.append(
                "章节「{0}」正文过短（仅 {1} 字），疑似没写后续/内容不完整。".format(title[:30], len(body.strip()))
            )


def _check_images(doc: Document, report: QualityReport) -> None:
    paras = list(_iter_body_paragraphs(doc))
    for i, para in enumerate(paras):
        if not _paragraph_has_drawing(para):
            continue
        nxt = ""
        for p in paras[i + 1:]:
            if p.text.strip():
                nxt = p.text.strip()
                break
        if nxt and not nxt.startswith(("图：", "图:", "表：", "表:")):
            report.warnings.append("图片后缺少图注：{0}…".format(nxt[:24]))
        elif not nxt:
            report.warnings.append("文档末尾的图片缺少图注。")


def _fix_markdown_residue(doc: Document, fixes: list) -> None:
    n = 0
    for para in list(_iter_body_paragraphs(doc)):
        text = para.text or ""
        if not text.strip():
            continue
        dirty = "**" in text or "`" in text or re.match(r"^\s*#{1,6}\s", text)
        if not dirty and (text.count("|") >= 2 and _heading_level(para) == 0):
            dirty = True
        if not dirty:
            continue
        cleaned = _strip_md_residue(text)
        if cleaned == text.strip():
            continue
        # 折叠为单 run，避免残留分散在多个 run 里
        first = para.runs[0] if para.runs else para.add_run()
        for run in list(para.runs):
            if run is not first:
                run._element.getparent().remove(run._element)
        first.text = cleaned
        n += 1
    if n:
        fixes.append("清理 {0} 个段落的 Markdown 残留标记".format(n))


def _fix_empty_paragraphs(doc: Document, fixes: list) -> None:
    paras = list(_iter_body_paragraphs(doc))
    run_start = 0
    for i, p in enumerate(paras[:4]):
        if not p.text.strip():
            run_start = i + 1
        else:
            break
    empty_run = []
    removed = 0
    for i in range(run_start, len(paras)):
        p = paras[i]
        if not p.text.strip():
            empty_run.append(p)
        else:
            # 只保留 1 个空段作为分隔
            if len(empty_run) > 1:
                for extra in empty_run[1:]:
                    extra._p.getparent().remove(extra._p)
                    removed += 1
            empty_run = []
    if len(empty_run) > 1:
        for extra in empty_run[1:]:
            extra._p.getparent().remove(extra._p)
            removed += 1
    if removed:
        fixes.append("移除 {0} 个多余空段落".format(removed))


def _fix_table_borders(doc: Document, fixes: list) -> None:
    n = 0
    for table in doc.tables:
        if not _table_has_borders(table):
            _ensure_table_borders(table)
            n += 1
    if n:
        fixes.append("为 {0} 个表格补上边框".format(n))


def _fix_body_fonts(doc: Document, fixes: list) -> None:
    """归一正文 run 字体：仅当 run 回退到常见默认字体时，对齐 Normal 样式。"""
    body_font = _normal_style_font(doc)
    if not body_font:
        return
    n = 0
    for para in _iter_body_paragraphs(doc):
        if _heading_level(para):
            continue
        for run in para.runs:
            f = _run_east_asia_font(run)
            if f is None:
                continue
            if f.lower() in FALLBACK_FONTS and f != body_font:
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = OxmlElement("w:rFonts")
                    rpr.append(rfonts)
                rfonts.set(qn("w:eastAsia"), body_font)
                n += 1
    if n:
        fixes.append("归一 {0} 个 run 的字体到正文样式「{1}」".format(n, body_font))


def verify_document(path) -> QualityReport:
    """只检查不修改，返回质量报告。"""
    report = QualityReport(path=str(path))
    try:
        doc = Document(str(path))
    except Exception as exc:  # noqa: BLE001
        report.issues.append("文档无法打开（可能已损坏）：{0}".format(exc))
        report.ok = False
        return report
    _check_headings(doc, report)
    _check_empty_paragraphs(doc, report)
    _check_md_residue(doc, report)
    _check_font_consistency(doc, report)
    _check_tables(doc, report)
    _check_section_content(doc, report)
    _check_images(doc, report)
    report.ok = not report.issues
    return report


def fix_document(path, log=None) -> QualityReport:
    """先自动修复安全项，再校验并返回报告。"""
    report = QualityReport(path=str(path))
    try:
        doc = Document(str(path))
    except Exception as exc:  # noqa: BLE001
        report.issues.append("文档无法打开（可能已损坏）：{0}".format(exc))
        report.ok = False
        return report

    _fix_markdown_residue(doc, report.fixes)
    _fix_empty_paragraphs(doc, report.fixes)
    _fix_table_borders(doc, report.fixes)
    _fix_body_fonts(doc, report.fixes)
    if report.fixes:
        try:
            doc.save(str(path))
        except Exception as exc:  # noqa: BLE001
            report.issues.append("保存修复结果失败：{0}".format(exc))

    # 修复后再完整校验一遍
    second = verify_document(path)
    report.issues = second.issues
    report.warnings = second.warnings
    report.ok = second.ok
    if log:
        for line in report.log_lines():
            log(line)
    return report
