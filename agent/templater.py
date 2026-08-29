"""模板填写模式：在给定的模板 docx 中直接填写内容。

底层逻辑（模板驱动，参考 docxtpl / report-forge / DocuFiller）：
1. analyze_template 解析模板 -> 蓝图（标题/空位/占位符/表格结构）
2. AI 只输出「哪里填什么」的结构化 JSON（anchor + content + rows），不输出排版指令
3. apply_fill 落笔：复制模板相邻段落的 rPr/pPr 格式、复制模板表格行格式，
   模板的页面/样式/表格/页眉页脚 100% 保留
"""
from __future__ import annotations

import re
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Optional

from docx import Document

from .config import Config
from .editor import (
    EditValidationError,
    _body_level_element,
    _insert_after_element,
    _normalize_text,
)
from .extractor import docx_to_markdown
from .llm import LLMClient, LLMError
from .memory import Memory

FILL_PROMPT = """你是一名文档填写助手。用户提供了一份 Word 模板（已转成结构化清单），以及一份写作指令。
你的任务：识别模板中【需要填写的位置】，并为这些位置生成内容。模板中已有内容的章节绝对不要动、不要改写。

【模板正文全文】里的每段文字都是模板内容：标了【要求】的语句（含 要求/注意/必须/格式/说明 等关键词）是模板作者写的要求，
必须逐条遵守（例如：格式要求、内容范围、字数要求、必须包含哪些内容）；未标【要求】的正文通常是模板已有内容，不要改写，
只需要在对应空位填写符合整体要求的内容。

模板清单里有三种可填位置，按优先级：
1. cells（最高优先级）：模板中的【表格填写位】和【正文标签填写位】，例如「实验目的及要求：」标签后空白、
   「实验项目｜（空）」「预习日期｜（空）」这类栏目。label 必须是清单里的标签原文（完全一致，用于定位）；
   value 为该位置要填的内容，多段用换行分隔。清单中标注为【信息栏】（姓名/学号/签名/成绩/日期/批阅等）的，
   除非用户指令明确要求，否则不填。清单中给出的【批注】是模板作者的要求（如格式、具体数据），必须遵守：
   批注里的具体数据（如“实验1：第1-3周；实验2：第4-11周；实验3：第12-14周；实验4：第15-16周”）要完整、逐项
   原样写入对应填写位，不要只取其中一部分，也不要自行简化。
2. sections：模板中【空白或没有实质正文】的章节标题（清单里标了【标题】的）。heading 必须是标题原文，
   content 为该章节正文，用 Markdown 编写：可用段落、`-` 无序列表、`1.` 有序列表；不要输出标题行、
   不要输出 markdown 表格、不要用 **加粗**、*斜体*、反引号、LaTeX 符号，单位用普通括号如 `L (m)`。
3. tables：仅当清单里有【数据表格】（表头行 + 空数据行）时，把数据行填进去；headers 必须与模板表头完全一致。
   【表单型表格】（标签+空白栏目）绝对不要追加任何新行，也不要改动其行结构。

内容充实度要求（很重要，必须遵守）：
- 实验步骤/实验内容/操作过程类填写位：用编号列出完整流程，不少于 4 步；每步写清具体操作、参数或命令、观察/预期结果，不要只写一两句。
- 实验目的/原理/结果分析/总结/心得体会类：内容要完整充实，一般 80~200 字，覆盖要点（目的写清目标与意义；原理写清依据与公式；分析写清数据与定量结论）。
- 数据记录类：示例数据至少 3 行，列名/单位与模板一致。
- 同一栏目有多个同名位置（如两张表都有「实验步骤」）时，按位置顺序分别给出对应内容（预习部分写预习步骤、实验部分写实验步骤），不要互相照抄。

没有真实数据时，用示例数据并让对应结论/分析明确说明数据为示例。
严格只输出 JSON，不要输出任何其他文字。格式：
{
  "cells": [
    {"label": "模板中的标签原文", "value": "该位置的内容（多段用\n分隔）"}
  ],
  "sections": [
    {"heading": "模板中的章节标题原文", "content": "该章节正文……"}
  ],
  "tables": [
    {"headers": ["列1", "列2"], "rows": [["值1", "值2"]]}
  ]
}"""

def _is_heading(para) -> bool:
    try:
        name = (para.style.name or "").strip()
    except Exception:
        return False
    low = name.lower()
    return low.startswith("heading") or low.startswith("标题")


_NUM_PREFIX_RE = re.compile(r"^[（(]?[0-9一二三四五六七八九十]+[、．.）)]?\s*")


def _match_score(anchor_text: str, target: str) -> int:
    """标题匹配打分：0=不匹配，分数越高越像同一标题。"""
    a = _normalize_text(anchor_text)
    t = _normalize_text(target)
    if not a or not t:
        return 0
    if a == t:
        return 100
    # 忽略「一、」这类序号前缀后再比（模板写「一、实验目的」，LLM 可能只回「实验目的」）
    a_no = _NUM_PREFIX_RE.sub("", a).strip()
    t_no = _NUM_PREFIX_RE.sub("", t).strip()
    if a_no and a_no == t_no:
        return 90
    if a_no and (a_no in t or t in a_no) and min(len(a_no), len(t)) >= 2:
        return 70
    return 0


def _looks_like_heading_para(para) -> bool:
    """判断段落是否像模板标题：Heading 样式，或 短文本 + 加粗/居中/独立行。"""
    if _is_heading(para):
        return True
    t = para.text.strip()
    if not t or len(t) > 40:
        return False
    try:
        bold = any(r.font.bold for r in para.runs)
        align = para.alignment
    except Exception:
        bold, align = False, None
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    centered = align == WD_ALIGN_PARAGRAPH.CENTER
    if bold or centered:
        return True
    # 序号开头 + 独立短行（如「1. 实验目的」「（一）概述」）
    import re as _re
    if _re.match(r"^[（(]?[0-9一二三四五六七八九十]+[、．.）)]", t):
        return True
    return False


def _find_heading_paragraph(doc: Document, text: str):
    """定位模板中的章节标题段落。

    现实模板的标题常是【普通段落+加粗/居中】（不是 Heading 样式），
    因此候选集 = 样式标题 + 形似标题的短段落；再按 原文精确->忽略序号->包含 打分。
    """
    norm = _normalize_text(text)
    if not norm:
        return None
    candidates = [p for p in doc.paragraphs if p.text.strip() and _looks_like_heading_para(p)]
    if not candidates:  # 兜底：找不到任何形似标题的段落时，退回全文匹配
        candidates = [p for p in doc.paragraphs if p.text.strip()]
    best = None
    best_score = 0
    for p in candidates:
        score = _match_score(p.text, text)
        if score > best_score:
            best_score = score
            best = p
    return best if best_score >= 70 else None


def _find_table_by_headers(doc: Document, headers: list[str]):
    """按表头序列定位模板表格：先精确匹配整行，再按列匹配打分。"""
    clean = [_normalize_text(h) for h in headers if _normalize_text(h)]
    if not clean:
        return None
    target_joined = "".join(clean)
    best_table = None
    best_score = 0
    for table in doc.tables:
        if not table.rows:
            continue
        cells = [_normalize_text(c.text) for c in table.rows[0].cells]
        if "".join(cells) == target_joined:
            return table
        score = sum(1 for c in cells if c and any(c == h or c in h or h in c for h in clean))
        if score > best_score:
            best_score = score
            best_table = table
    return best_table if best_score >= max(1, len(clean) * 0.6) else None


def _append_table_rows(table, rows: list[list[str]]) -> int:
    """向模板表格追加数据行，行格式继承模板已有行。"""
    template_run = None
    if len(table.rows) >= 2:  # 用最后一行做格式模板（表头除外）
        for cell in table.rows[-1].cells:
            for p in cell.paragraphs:
                if p.runs:
                    template_run = p.runs[0]
                    break
            if template_run is not None:
                break
    added = 0
    for row in rows:
        values = [str(v) for v in (row or [])]
        if not any(v.strip() for v in values):
            continue
        cells = table.add_row().cells
        for j, val in enumerate(values):
            if j >= len(cells):
                break
            p = cells[j].paragraphs[0]
            if p.runs:
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
            run = p.add_run(val.strip())
            if template_run is not None and template_run._element.rPr is not None:
                run._element.insert(0, deepcopy(template_run._element.rPr))
        added += 1
    return added


def plan_fill(command: str, blueprint: str, llm: LLMClient) -> dict[str, Any]:
    """AI 基于模板蓝图，输出 {sections, tables, cells} 填写计划（纯内容，不含格式）。"""
    messages = [
        {"role": "system", "content": FILL_PROMPT},
        {
            "role": "user",
            "content": f"【用户写作指令】\n{command}\n\n【模板结构清单】\n{blueprint}",
        },
    ]
    error = ""
    try:
        data = llm.chat_json(messages, temperature=0.2, max_tokens=8192)
    except LLMError as exc:
        data = {}
        error = str(exc)
    if not isinstance(data, dict):
        data = {}
    sections = []
    for sec in data.get("sections") or []:
        if not isinstance(sec, dict):
            continue
        anchor = str(sec.get("anchor") or sec.get("heading") or "").strip()
        content = str(sec.get("content", "")).strip()
        if anchor and content:
            sections.append({"anchor": anchor, "content": content})
    tables = []
    for t in data.get("tables") or []:
        if not isinstance(t, dict):
            continue
        headers = [str(h).strip() for h in (t.get("headers") or []) if str(h).strip()]
        rows = [[str(v).strip() for v in (r or [])] for r in (t.get("rows") or [])]
        if headers:
            tables.append({"headers": headers, "rows": rows})
    cells = []
    for c in data.get("cells") or []:
        if not isinstance(c, dict):
            continue
        ph = str(c.get("label") or c.get("placeholder") or c.get("name") or "").strip()
        val = str(c.get("value") or c.get("content") or "").strip()
        if ph and val:
            cells.append({"label": ph, "value": val})
    return {"sections": sections, "tables": tables, "cells": cells, "error": error}


def apply_fills(doc: Document, fill: dict, log: Callable[[str], None] = print) -> list[str]:
    """把填写计划写入模板 docx。"""
    applied: list[str] = []
    for sec in fill.get("sections") or []:
        heading = sec["heading"]
        anchor = _find_heading_paragraph(doc, heading)
        if anchor is None:
            log(f"   ⚠ 未找到章节标题「{heading[:30]}」，已跳过（标题必须与模板原文一致）")
            continue
        _insert_after_element(doc, _body_level_element(doc, anchor), sec["content"])
        applied.append(f"填写章节「{heading[:30]}」")
    for t in fill.get("tables") or []:
        table = _find_table_by_headers(doc, t["headers"])
        if table is None:
            log(f"   ⚠ 未找到表头匹配的表格「{'、'.join(t['headers'])[:40]}」，已跳过")
            continue
        n = _append_table_rows(table, t["rows"])
        if n:
            applied.append(f"填写表格「{'、'.join(t['headers'])[:30]}」（{n} 行数据）")
        else:
            log(f"   ⚠ 表格「{'、'.join(t['headers'])[:30]}」没有有效数据行，已跳过")
    return applied


def fill_template(
    template_path: Path,
    command: str,
    config: Config,
    memory: Memory,
    llm: Optional[LLMClient] = None,
    log: Callable[[str], None] = print,
) -> Path:
    """在给定模板 docx 中填写内容，保留模板全部格式，另存为新文件。"""
    memory.add_user(command)
    llm = llm or LLMClient(config)
    template_path = Path(template_path)
    if not template_path.is_file():
        raise FileNotFoundError(f"模板文件不存在：{template_path}")

    from .template_engine import analyze_template, blueprint_text, apply_fill as engine_apply

    doc = Document(str(template_path))
    info = analyze_template(template_path, doc=doc)
    log(f"① 解析模板：{template_path.name}（{len(info['blocks'])} 个位置、{len(info['tables'])} 个表格）")

    log("② AI 分析模板待填位置并生成内容 ...")
    fill = plan_fill(command, blueprint_text(info), llm)
    if fill.get("error"):
        log(f"   ⚠ AI 分析模板失败：{fill['error']}")
    log(f"   计划填写 {len(fill['sections'])} 个章节、{len(fill['tables'])} 个表格、{len(fill['cells'])} 个占位符")

    log("③ 写入模板（格式继承模板段落/表格）...")
    work_path = Path(str(template_path) + ".filling.docx")
    import shutil
    shutil.copy(template_path, work_path)
    report: dict = {}
    applied = engine_apply(work_path, fill, log=log, report=report)
    for item in applied:
        log(f"   ✓ {item}")
    if not applied:
        log("⚠ 没有可填写的位置：模板可能已完整，或章节标题/表头未匹配上。")

    out_dir = Path(getattr(config, "output_dir_override", None) or config.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = out_dir / f"{template_path.stem}_填写_{stamp}.docx"
    shutil.move(str(work_path), str(out_path))

    check = Document(str(out_path))
    if not check.paragraphs and not check.tables:
        raise EditValidationError("保存后的文档为空，已中止。")
    log("   ✓ 完整性校验通过")

    _log_fill_report(report, log)
    _save_fill_report(out_path, report, log)

    memory.add_result(
        {
            "action": "fill",
            "title": template_path.name,
            "file": str(out_path),
            "sections": len(fill["sections"]),
            "tables": len(fill["tables"]),
        }
    )
    if hasattr(llm, "usage_text"):
        log(f"   ⚙ {llm.usage_text()}")
    log(f"✔ 填写完成：{out_path}")
    return out_path

def _log_fill_report(report: dict, log: Callable[[str], None] = print) -> None:
    """把填写完成度写入日志，跳过的位置给出原因，方便用户核对。"""
    cells_ok = report.get("cells_ok", 0)
    cells_skip = report.get("cells_skip", [])
    sections_ok = report.get("sections_ok", 0)
    sections_skip = report.get("sections_skip", [])
    tables_ok = report.get("tables_ok", 0)
    tables_skip = report.get("tables_skip", [])
    log(
        f"   📋 填写报告：栏目 {cells_ok} 处成功"
        + (f"、{len(cells_skip)} 处跳过" if cells_skip else "")
        + f"；章节 {sections_ok} 节成功"
        + (f"、{len(sections_skip)} 节跳过" if sections_skip else "")
        + f"；表格 {tables_ok} 行数据"
    )
    for item in cells_skip + sections_skip + tables_skip:
        log(f"   ⚠ 跳过：{item}")
    for item in report.get("cells_thin") or []:
        log(f"   ⚠ 偏短：{item}")
    for item in report.get("cells_note") or []:
        log(f"   ℹ {item}")


def _save_fill_report(out_path: Path, report: dict, log: Callable[[str], None] = print) -> None:
    """有跳过项时，在输出同目录保存一份可读的填写报告（.txt），方便留存与核对。"""
    thin = (report.get("cells_thin") or []) + (report.get("cells_note") or [])
    skip = (report.get("cells_skip") or []) + (report.get("sections_skip") or []) + (report.get("tables_skip") or [])
    if not skip and not thin:
        return
    lines = [
        "WordAgent 模板填写报告",
        "生成时间：" + datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        f"输出文档：{out_path}",
        f"成功：栏目 {report.get('cells_ok', 0)}、章节 {report.get('sections_ok', 0)}、表格行 {report.get('tables_ok', 0)}",
        f"跳过：{len(skip)} 项",
    ]
    if thin:
        lines.append(f"提示：{len(thin)} 项（内容偏短 / 补齐说明）")
    lines.append("")
    if skip:
        lines.append("跳过明细（请人工核对是否需要在 Word 中补充）：")
        lines += [f"- {x}" for x in skip]
    if thin:
        lines.append("")
        lines.append("提示明细（内容偏短或同名补齐，可在 Word 中补充完善）：")
        lines += [f"- {x}" for x in thin]
    try:
        rp = out_path.with_suffix(".填写报告.txt")
        rp.write_text("\n".join(lines), encoding="utf-8-sig")
        log(f"   📄 填写报告已保存：{rp}")
    except OSError:
        pass
