"""docx -> 结构化 Markdown 提取（编辑模式用，保留段落/表格/样式信息）。"""
from __future__ import annotations

from dataclasses import dataclass
from typing import Iterator, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
import re
from docx.text.paragraph import Paragraph


@dataclass
class ParaInfo:
    index: int          # 序号（供提示词引用）
    kind: str           # paragraph | cell | image | page_break
    style: str          # 样式名（Heading 1 / Normal / List Bullet ...）
    text: str           # 提取出的文本（含 markdown 行内标记）
    para: Optional[Paragraph] = None  # 对应 python-docx 段落对象


def _iter_blocks(doc: Document) -> Iterator[Paragraph | Table]:
    """按文档顺序遍历正文中的段落与表格（含表格内段落）。"""
    body = doc.element.body

    def walk(element):
        for child in element.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, doc)
            elif child.tag == qn("w:tbl"):
                yield Table(child, doc)
                yield from walk(child)

    yield from walk(body)


def _heading_level(para: Paragraph) -> Optional[int]:
    style = para.style
    name = (style.name if style is not None else "") or ""
    low = name.lower()
    # 只认标题类样式（Heading N / 标题 N），避免 "List Bullet 2" 等被误判为标题
    if low.startswith(("heading", "标题")):
        m = re.search(r"\d+", name)
        return int(m.group()) if m else 1
    ppr = para._p.pPr
    if ppr is not None:
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is not None:
            try:
                return int(outline.get(qn("w:val"), "0")) + 1
            except ValueError:
                return 1
    return None


def _style_name(para: Paragraph) -> str:
    try:
        return para.style.name if para.style is not None else "Normal"
    except Exception:
        return "Normal"


def _runs_to_markdown(para: Paragraph) -> str:
    if not para.runs:
        return para.text
    parts: list[str] = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.font.bold:
            text = f"**{text}**"
        if run.font.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts) if parts else para.text


def _para_to_markdown(para: Paragraph, index: int) -> tuple[Optional[str], ParaInfo]:
    style_name = _style_name(para)
    level = _heading_level(para)
    low = style_name.lower()
    text = _runs_to_markdown(para)

    # 图片占位
    if para._p.find(qn("w:drawing")) is not None or para._p.find(qn("w:pict")) is not None:
        return f"[图片 {index}]", ParaInfo(index, "image", style_name, f"[图片 {index}]", para)

    # 分页符
    has_page_break = any(
        br.get(qn("w:type")) == "page"
        for br in para._p.iter(qn("w:br"))
    )
    if has_page_break:
        return "[分页符]", ParaInfo(index, "page_break", style_name, "[分页符]", para)

    if level:
        prefix = "#" * min(level, 4)
        markdown = f"{prefix} {text}" if text else f"{prefix} "
    elif low.startswith("list bullet"):
        markdown = f"- {text}" if text else "-"
    elif low.startswith("list number"):
        markdown = f"1. {text}" if text else "1."
    elif low.startswith(("quote", "intense quote")):
        markdown = f"> {text}" if text else ">"
    else:
        markdown = text
    return markdown, ParaInfo(index, "paragraph", style_name, text, para)


def _table_to_markdown(table: Table) -> list[str]:
    rows: list[str] = []
    for row_index, row in enumerate(table.rows):
        # 合并单元格会重复出现在多列：按底层 _tc 去重，避免 AI 看到重复列
        seen: set[int] = set()
        cells: list[str] = []
        for cell in row.cells:
            key = id(cell._tc)
            if key in seen:
                continue
            seen.add(key)
            cells.append(cell.text.replace("\n", " ").strip())
        rows.append("| " + " | ".join(cells) + " |")
        if row_index == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return rows


def doc_requirements(path: str) -> str:
    """提取文档内要求文本：批注 + 页眉页脚 + 文本框/形状文字。

    用户的重要文件里，很多要求写在批注、页眉页脚或文本框里，编辑前必须读到。
    返回多行文本，无内容时返回空串。
    """
    import zipfile
    import re as _re
    lines: list[str] = []
    doc = Document(str(path))
    # 1) 批注（comments.xml）
    comments: dict[str, str] = {}
    try:
        with zipfile.ZipFile(str(path)) as zf:
            if "word/comments.xml" in zf.namelist():
                cxml = zf.read("word/comments.xml").decode("utf-8", "ignore")
                from lxml import etree as _et
                croot = _et.fromstring(cxml.encode("utf-8"))
                for c in croot.iter():
                    if c.tag == qn("w:comment"):
                        cid = c.get(qn("w:id"))
                        txt = "".join(t.text or "" for t in c.iter(qn("w:t")))
                        txt = _re.sub(r"\s+", " ", txt).strip()
                        if cid and txt:
                            comments[cid] = txt
    except Exception:
        pass
    if comments:
        # 按文档顺序输出，附所在段落文本
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        body = doc.element.body

        def walk(el) -> None:
            for child in el.iterchildren():
                if child.tag == qn("w:p"):
                    para = Paragraph(child, doc)
                    ids = [c.get(qn("w:id")) for c in child.iter(qn("w:commentRangeStart"))]
                    txts = [comments.get(i, "") for i in ids if comments.get(i)]
                    if txts:
                        ctx = para.text.strip()[:60]
                        lines.append(f"【批注】{ '；'.join(txts) }" + (f"（位置：{ctx}）" if ctx else ""))
                elif child.tag == qn("w:tbl"):
                    tbl = Table(child, doc)
                    for row in tbl.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                ids = [c.get(qn("w:id")) for c in para._p.iter(qn("w:commentRangeStart"))]
                                txts = [comments.get(i, "") for i in ids if comments.get(i)]
                                if txts:
                                    ctx = para.text.strip()[:60]
                                    lines.append(f"【批注】{ '；'.join(txts) }" + (f"（位置：{ctx}）" if ctx else ""))
                    walk(child)

        walk(body)
    # 2) 页眉页脚
    try:
        for sec in doc.sections:
            for p_ in sec.header.paragraphs:
                t = p_.text.strip()
                if t:
                    lines.append("【页眉】" + t)
            for p_ in sec.footer.paragraphs:
                t = p_.text.strip()
                if t:
                    lines.append("【页脚】" + t)
    except Exception:
        pass
    # 3) 文本框 / 形状内文字（w:txbxContent）
    try:
        for txbx in doc.element.body.iter(qn("w:txbxContent")):
            texts = [t.text or "" for t in txbx.iter(qn("w:t"))]
            t = "".join(texts).strip()
            if t:
                lines.append("【文本框】" + t)
    except Exception:
        pass
    seen, out = set(), []
    for ln in lines:
        if ln not in seen:
            seen.add(ln)
            out.append(ln)
    return "\n".join(out)


def docx_to_markdown(doc: Document) -> tuple[str, list[ParaInfo]]:
    """提取文档为 Markdown 文本，同时返回段落索引（含对象引用，供编辑落笔）。"""
    lines: list[str] = []
    index: list[ParaInfo] = []
    counter = 0
    for block in _iter_blocks(doc):
        if isinstance(block, Paragraph):
            counter += 1
            markdown, info = _para_to_markdown(block, counter)
            index.append(info)
            if markdown is not None:
                lines.append(markdown)
        else:  # Table
            rows = _table_to_markdown(block)
            lines.extend(rows)
            for row in block.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        counter += 1
                        index.append(ParaInfo(counter, "cell", _style_name(para), para.text, para))
    return "\n".join(lines), index
