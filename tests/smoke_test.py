"""离线冒烟测试：不调用真实 API，验证 生成/编辑 全流程与质量机制。

运行：python tests/smoke_test.py

覆盖：
- 生成模式全流程（规划/分节/排版/表格/记忆）
- 编辑：三级匹配（精确/包含/近似）容错标点与加粗标记
- 编辑：歧义目标 -> AmbiguousTargetError / 人工选择
- 编辑：AI 复核循环（发现遗漏 -> 补充操作）
- 编辑：表格行锚定插入（新章节落到章节末尾）
- 编辑：插入位置正确性（不打断子节 / 子级标题紧邻插入）
- 编辑：另存新文件 + 备份；覆盖原文件 + 备份
- 编辑：目标不存在 -> 中止且不写盘
"""
from __future__ import annotations

import json
import shutil
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document  # noqa: E402

from agent.config import Config  # noqa: E402
from agent.editor import AmbiguousTargetError, EditValidationError, edit_document  # noqa: E402
from agent.llm import LLMClient  # noqa: E402
from agent.memory import Memory  # noqa: E402
from agent.pipeline import run_pipeline, safe_filename  # noqa: E402
from agent.refdocs import collect_reference_context  # noqa: E402
from agent.renderer import render_markdown_to_docx  # noqa: E402

SOURCE_MD = (
    "## 项目背景\n\n这是背景段落。\n\n## 核心功能\n\n"
    "- 功能点A\n- 功能点B\n- 删除我\n\n## 总结\n\n结束语。"
)


class FakeGenerateLLM(LLMClient):
    """模拟生成模式：规划给 JSON，写作给 Markdown。"""

    def __init__(self):
        self.config = Config(api_key="fake-key")

    def chat(self, messages, temperature=None, max_tokens=None, json_mode=False):
        if json_mode:
            return (
                '{"title":"智能办公软件产品需求文档","style":"business","language":"zh-CN",'
                '"toc":true,"sections":['
                '{"heading":"项目背景","level":1,"description":"行业背景与痛点"},'
                '{"heading":"核心功能","level":1,"description":"功能清单"}]}'
            )
        return (
            "## 项目背景\n\n随着远程办公普及，企业需要更**高效**的协作工具。\n\n"
            "## 核心功能\n\n1. 文档在线编辑\n2. 多人协同批注\n\n"
            "| 功能 | 优先级 |\n| --- | --- |\n| 在线编辑 | P0 |\n| 协同批注 | P1 |"
        )


class FakeEditLLM(LLMClient):
    """模拟编辑模式。verify_satisfied=False 时第一轮复核返回补充操作。"""

    def __init__(self, verify_satisfied: bool = True, extra_ops: list[dict] | None = None,
                 plan_json: str | None = None):
        self.config = Config(api_key="fake-key")
        self.verify_satisfied = verify_satisfied
        self.extra_ops = extra_ops or []
        self.plan_json = plan_json
        self.critic_calls = 0

    def chat(self, messages, temperature=None, max_tokens=None, json_mode=False):
        system = str(messages[0].get("content", ""))
        if json_mode:
            if "质量检查员" in system:  # AI 复核
                self.critic_calls += 1
                if self.verify_satisfied or self.critic_calls > 1:
                    return '{"satisfied": true, "reason": "ok", "operations": []}'
                return (
                    '{"satisfied": false, "reason": "补充修改", "operations": '
                    + json.dumps(self.extra_ops, ensure_ascii=False) + "}"
                )
            if self.plan_json is not None:
                return self.plan_json
            # 默认计划：target 带加粗标记与多余标点，验证容错匹配
            return (
                '{"summary":"测试修改：改背景、插功能、删多余项",'
                '"operations":['
                '{"op":"replace","target":"这是**背景**段落。","new_text":"这是**修改后**的背景段落"},'
                '{"op":"insert_after","target":"功能点A","new_text":"- 新增功能点C"},'
                '{"op":"delete","target":"删除我"}]}'
            )
        raise AssertionError("编辑模式不应触发普通 chat")


def _make_source_doc(path: Path) -> None:
    render_markdown_to_docx(SOURCE_MD, path, title="测试原文档", style="business", toc=False)


def _doc_text(path: Path) -> str:
    return "\n".join(p.text for p in Document(str(path)).paragraphs)


def _test_generate(tmp: Path) -> None:
    config = Config(api_key="fake-key", output_dir=tmp / "out", memory_file=tmp / "memory.json")
    memory = Memory(config.memory_file)
    llm = FakeGenerateLLM()
    path = run_pipeline("写一份产品需求文档", config, memory, llm=llm)
    assert path.exists() and path.suffix == ".docx"
    doc = Document(str(path))
    assert len(doc.tables) >= 1 and "项目背景" in "\n".join(p.text for p in doc.paragraphs)
    assert len(memory.entries) == 2
    assert safe_filename('a/b:c*d?"e<f>g|h') == "abcdefgh"
    print("✔ 生成模式通过")


def _test_edit_quality_matching(tmp: Path) -> None:
    src = tmp / "质量匹配.docx"
    _make_source_doc(src)
    config = Config(api_key="fake-key", output_dir=tmp / "out", memory_file=tmp / "memory.json")
    memory = Memory(config.memory_file)
    out = edit_document(src, "修改一下", config, memory, llm=FakeEditLLM(), output_dir=tmp / "out")
    assert out is not None
    text = _doc_text(out)
    # 带 ** 标记 + 多余标点的 target 仍能命中（归一化 + 模糊匹配）
    assert "修改后" in text and "新增功能点C" in text and "删除我" not in text
    # 原文件未动
    assert "删除我" in _doc_text(src)
    print("✔ 编辑容错匹配（标记/标点差异）通过")


def _test_edit_ambiguous(tmp: Path) -> None:
    src = tmp / "歧义.docx"
    _make_source_doc(src)
    config = Config(api_key="fake-key", output_dir=tmp / "out", memory_file=tmp / "memory.json")
    memory = Memory(config.memory_file)

    # 1) 无选择器 -> 抛 AmbiguousTargetError，不写盘
    ambiguous_llm = FakeEditLLM(plan_json='{"summary":"歧义","operations":[{"op":"delete","target":"功能点"}]}')
    before = set(tmp.glob("out/*.docx")) | set((tmp / "out" / "backups").glob("*.docx"))
    try:
        edit_document(src, "删掉功能点", config, memory, llm=ambiguous_llm, output_dir=tmp / "out")
        raise AssertionError("应抛出 AmbiguousTargetError")
    except AmbiguousTargetError as exc:
        assert len(exc.candidates) >= 2
    after = set(tmp.glob("out/*.docx")) | set((tmp / "out" / "backups").glob("*.docx"))
    assert after == before, "歧义未解决时不应写盘"

    # 2) 提供选择器 -> 删除被选中的候选
    memory = Memory(config.memory_file)  # 重置记忆
    def pick(op, candidates):
        return candidates[0]
    out = edit_document(src, "删掉功能点", config, memory, llm=ambiguous_llm,
                        output_dir=tmp / "out", resolve_ambiguous=pick)
    assert out is not None
    text = _doc_text(out)
    assert "功能点A" not in text and "功能点B" in text
    print("✔ 编辑歧义处理（中止不写盘 / 人工选择）通过")


def _test_edit_verify_loop(tmp: Path) -> None:
    src = tmp / "复核.docx"
    _make_source_doc(src)
    config = Config(api_key="fake-key", output_dir=tmp / "out", memory_file=tmp / "memory.json")

    # 复核发现遗漏（功能点B 也要删）-> 自动补充执行
    memory = Memory(config.memory_file)
    llm = FakeEditLLM(verify_satisfied=False, extra_ops=[{"op": "delete", "target": "功能点A"}, {"op": "delete", "target": "功能点B"}])
    out = edit_document(src, "删除所有功能点", config, memory, llm=llm, output_dir=tmp / "out", auto_fix=True)
    assert out is not None
    text = _doc_text(out)
    assert "功能点A" not in text and "功能点B" not in text, "复核应补齐遗漏操作"
    assert llm.critic_calls >= 2

    # 关闭复核 -> 不补齐
    memory = Memory(config.memory_file)
    llm2 = FakeEditLLM(verify_satisfied=False, extra_ops=[{"op": "delete", "target": "功能点B"}])
    out2 = edit_document(src, "删除所有功能点", config, memory, llm=llm2, output_dir=tmp / "out", verify=False)
    text2 = _doc_text(out2)
    assert "功能点A" in text2 and "功能点B" in text2, "关闭复核时不应补齐"
    print("✔ AI 复核循环（补齐遗漏 / 可关闭）通过")


def _test_edit_table_insert(tmp: Path) -> None:
    """表格行作为锚定点：插入内容若以标题开头，自动放到表格所在章节的末尾。"""
    src = tmp / "表格文档.docx"
    render_markdown_to_docx(
        "## 预算\n\n| 项目 | 金额 |\n| --- | --- |\n| 宣讲会 | 30,000 |\n| 线上广告 | 20,000 |\n\n## 总结\n\n结束语。",
        src, title="表格测试", style="business", toc=False,
    )
    class TableLLM(FakeEditLLM):
        def __init__(self):
            super().__init__()
            self.plan_json = (
                '{"summary":"在表格后新增内容",'
                '"operations":[{"op":"insert_after","target":"| 线上广告 | 20,000 |",'
                '"new_text":"## 新增节\\n- 项A\\n- 项B"}]}'
            )
    config = Config(api_key="fake-key", output_dir=tmp / "out", memory_file=tmp / "memory.json")
    memory = Memory(config.memory_file)
    out = edit_document(src, "在表格后新增内容", config, memory,
                        llm=TableLLM(), output_dir=tmp / "out")
    assert out is not None
    doc = Document(str(out))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    # 新增章节在表格所在章节（预算）末尾、总结之前
    idx_head = paras.index("新增节")
    idx_sum = paras.index("总结")
    assert idx_head < idx_sum, "新增节应在总结前"
    assert "项A" in paras and "项B" in paras
    print("✔ 编辑表格行键定（插入到表格所在章节末尾）通过")


def _test_edit_insert_position(tmp: Path) -> None:
    """新增章节锚定到子节内的表格行/正文段时，必须落到章节末尾，不得打断子节。"""
    src = tmp / "插入位置.docx"
    render_markdown_to_docx(
        "## 招聘预算\n\n### 一、总体预算\n内容一\n\n### 二、渠道费用\n内容二\n\n"
        "### 三、场地费用预算\n| 项目 | 金额 |\n| --- | --- |\n| 企业开放日 | 50,000 |\n| 宣讲会 | 30,000 |\n\n"
        "### 四、宣传物料\n内容四\n\n### 五、人力成本\n内容五\n\n## 入职与培训\n入职培训安排。",
        src, title="插入位置测试", style="business", toc=False,
    )

    class TableRowLLM(FakeEditLLM):
        def __init__(self):
            super().__init__()
            self.plan_json = (
                '{"summary":"新增风险应对预案章节",'
                '"operations":[{"op":"insert_after","target":"| 企业开放日 | 50,000 |",'
                '"new_text":"## 风险应对预案\\n- 预案A\\n- 预案B"}]}'
            )

    config = Config(api_key="fake-key", output_dir=tmp / "out", memory_file=tmp / "memory.json")
    memory = Memory(config.memory_file)
    out = edit_document(src, "新增风险应对预案章节", config, memory,
                        llm=TableRowLLM(), output_dir=tmp / "out")
    assert out is not None
    doc = Document(str(out))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    idx_new = paras.index("风险应对预案")
    idx_next_chapter = paras.index("入职与培训")
    # 新章节在所有 ### 子节内容之后、下一章之前
    assert idx_new > paras.index("内容五"), "新章节不应插在子节中间"
    assert idx_new < idx_next_chapter, "新章节应位于下一个 ## 章节之前"
    assert "内容四" in paras and "内容五" in paras

    # 子级标题（###）插在正文段落后仍保持紧邻（不位移）
    class SubLLM(FakeEditLLM):
        def __init__(self):
            super().__init__()
            self.plan_json = (
                '{"summary":"在背景段落后加补充小节",'
                '"operations":[{"op":"insert_after","target":"内容一",'
                '"new_text":"### 补充细节\\n- 细节A"}]}'
            )

    memory = Memory(config.memory_file)
    out2 = edit_document(src, "在背景段落后加补充小节", config, memory,
                         llm=SubLLM(), output_dir=tmp / "out")
    assert out2 is not None
    paras2 = [pp.text for pp in Document(str(out2)).paragraphs if pp.text.strip()]
    assert paras2.index("补充细节") == paras2.index("内容一") + 1, "### 子节应紧邻锚点段落"
    assert paras2.index("细节A") < paras2.index("内容二")
    print("✔ 插入位置正确性（新章节不打断子节 / 子级标题紧邻插入）通过")


def _test_reference_files(tmp: Path) -> None:
    """参考文件：csv 文本注入上下文、图片嵌入 docx。"""
    import base64
    png_1px = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
    )
    data = tmp / "shiyan_data.csv"
    data.write_text("组别,测量值(mm),理论值(mm)\n1,10.2,10.0\n2,9.8,10.0\n3,10.1,10.0\n", encoding="utf-8")
    img = tmp / "shiyan_screenshot.png"
    img.write_bytes(png_1px)

    refs = collect_reference_context([data, img])
    assert "10.2" in refs["writer_context"], "csv 数据应注入写作上下文"
    assert str(img.resolve()) in refs["images"], "图片路径应被收集"
    assert "shiyan_data.csv" in refs["planner_summary"], "文件名应出现在规划摘要"

    out = tmp / "with_image.docx"
    render_markdown_to_docx(
        "## 数据记录\n\n| 组别 | 测量值 |\n| --- | --- |\n| 1 | 10.2 |\n\n"
        "## 结果分析\n\n![实验过程截图](%s)\n\n结论。" % img.resolve(),
        out, title="带图测试", style="report", toc=False,
        extra_images=[str(img.resolve())],
    )
    doc = Document(str(out))
    assert len(doc.inline_shapes) >= 1, "docx 中应包含嵌入图片"
    print("✔ 参考文件（csv 文本 / 图片嵌入）通过")


def _test_edit_safety(tmp: Path) -> None:
    src = tmp / "安全.docx"
    _make_source_doc(src)
    config = Config(api_key="fake-key", output_dir=tmp / "out", memory_file=tmp / "memory.json")

    # 覆盖模式 + 备份
    memory = Memory(config.memory_file)
    out = edit_document(src, "修改一下", config, memory, llm=FakeEditLLM(),
                        output_dir=tmp / "out", save_as_new=False)
    assert out == src and "修改后" in _doc_text(src)
    assert len(list((tmp / "out" / "backups").glob("*.docx"))) >= 1

    # 目标不存在 -> 中止不写盘
    class BadLLM(FakeEditLLM):
        def __init__(self):
            super().__init__()
            self.plan_json = '{"summary":"非法","operations":[{"op":"replace","target":"完全不存在的文本XYZ","new_text":"x"}]}'

    memory = Memory(config.memory_file)
    before = set(tmp.glob("out/*.docx")) | set((tmp / "out" / "backups").glob("*.docx"))
    try:
        edit_document(src, "改不存在内容", config, memory, llm=BadLLM(), output_dir=tmp / "out")
        raise AssertionError("应抛出 EditValidationError")
    except EditValidationError:
        pass
    after = set(tmp.glob("out/*.docx")) | set((tmp / "out" / "backups").glob("*.docx"))
    assert after == before
    print("✔ 编辑安全（覆盖备份 / 失败不写盘）通过")
 
 
def _test_quality_gate(tmp: Path) -> None:
    """质量门禁：能发现并自动修复 Markdown 残留/空段污染/无边框表格。"""
    from agent.verify import _table_has_borders, fix_document, verify_document

    out = tmp / "质量门禁.docx"
    doc = Document()
    doc.add_paragraph()  # 标题页留白
    doc.add_paragraph()
    doc.add_heading("实验报告", level=1)
    doc.add_heading("数据记录", level=2)
    p = doc.add_paragraph()
    p.add_run("这是**加粗残留**和`代码残留`段落")
    for _ in range(4):  # 连续空段污染
        doc.add_paragraph()
    tbl = doc.add_table(rows=2, cols=2)  # 无边框表格
    tbl.cell(0, 0).text = "A"
    tbl.cell(1, 1).text = "B"
    doc.add_paragraph("结果分析")
    doc.save(str(out))

    report = verify_document(out)
    assert report.issues, "质量门禁应发现文档问题"
    assert any("残留" in i or "Markdown" in i for i in report.issues)

    report2 = fix_document(out)
    assert report2.fixes, "质量门禁应执行自动修复"
    doc2 = Document(str(out))
    text_all = "\n".join(x.text for x in doc2.paragraphs)
    assert "**加粗残留**" not in text_all and "`代码残留`" not in text_all
    empties = sum(1 for para in doc2.paragraphs if not para.text.strip())
    assert empties <= 3, "连续空段应被压缩（标题页留白 + 1 个分隔）"
    assert _table_has_borders(doc2.tables[0]), "无边框表格应补上边框"
    print("✔ 质量门禁（校验/自动修复）通过")


def main() -> int:
    tmp = Path(tempfile.mkdtemp(prefix="wordagent_test_"))
    try:
        _test_generate(tmp)
        _test_reference_files(tmp)
        _test_edit_quality_matching(tmp)
        _test_edit_ambiguous(tmp)
        _test_edit_verify_loop(tmp)
        _test_edit_table_insert(tmp)
        _test_edit_insert_position(tmp)
        _test_edit_safety(tmp)
        _test_quality_gate(tmp)
        print("\n✔ 全部冒烟测试通过")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)
    return 0


if __name__ == "__main__":
    sys.exit(main())
